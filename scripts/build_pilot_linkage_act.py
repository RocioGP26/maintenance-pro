"""Genera el modelo Word del Acta de Vinculacion al Piloto de Roustix."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "production-readiness" / "templates" / "ACTA-VINCULACION-PILOTO-ROUSTIX.docx"
CORPORATE_LOGO = ROOT / "static" / "img" / "roustix-logo-docx.png"

NAVY = "0A2540"
BLUE = "185FA5"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F3F7FB"
GRAY = "5B6472"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D8DEE7"
WHITE = "FFFFFF"
BLACK = "1D2733"
GREEN = "1F6B45"
GOLD = "8A6400"
RED = "9B1C1C"
USABLE_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths: list[int], indent: int = TABLE_INDENT_DXA) -> None:
    if sum(widths) != USABLE_DXA:
        raise ValueError(f"Anchos invalidos: {sum(widths)} != {USABLE_DXA}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(USABLE_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            cell.width = Twips(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, *, size=None, color=BLACK, bold=None, italic=None) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_cell_text(cell, text: str, *, bold=False, color=BLACK, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)


def style_table(table, *, header=True, label_columns: set[int] | None = None) -> None:
    label_columns = label_columns or set()
    table.style = "Table Grid"
    for r_idx, row in enumerate(table.rows):
        set_keep_together(row)
        if r_idx == 0 and header:
            set_repeat_table_header(row)
        for c_idx, cell in enumerate(row.cells):
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r_idx == 0 and header:
                set_cell_shading(cell, BLUE)
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_run_font(run, size=9.5, color=WHITE, bold=True)
            elif c_idx in label_columns:
                set_cell_shading(cell, LIGHT_BLUE)
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_run_font(run, size=9.5, color=NAVY, bold=True)
            else:
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_run_font(run, size=9.5, color=BLACK)


def add_label_detail_table(doc, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    for label, value in rows:
        cells = table.add_row().cells
        add_cell_text(cells[0], label, bold=True, color=NAVY)
        add_cell_text(cells[1], value)
    set_table_geometry(table, [2700, 6660])
    style_table(table, header=False, label_columns={0})
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_heading(doc, text: str, level=1) -> None:
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True


def add_field_line(doc, label: str, value: str = "____________________________________________________________") -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(f"{label}: ")
    set_run_font(r, size=10.5, color=NAVY, bold=True)
    r = p.add_run(value)
    set_run_font(r, size=10.5, color=BLACK)


def add_checkbox(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.first_line_indent = Inches(-0.02)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    marker = p.add_run("☐  ")
    set_run_font(marker, size=11, color=BLUE, bold=True)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=BLACK)


def add_callout(doc, title: str, body: str, *, fill=PALE_BLUE, accent=BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=accent, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(body)
    set_run_font(r, size=10, color=BLACK)
    set_table_geometry(table, [USABLE_DXA])
    style_table(table, header=False)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_word_field(paragraph, instruction: str, fallback: str = "1") -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = fallback
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])
    set_run_font(run, size=8.5, color=GRAY)


def enable_field_updates(doc: Document) -> None:
    """Pide a Word actualizar PAGE y NUMPAGES al abrir o imprimir."""
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, NAVY, 10, 5),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("ROUSTIX  ·  PILOTO CONTROLADO")
    set_run_font(r, size=8.5, color=GRAY, bold=True)

    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    left, right = table.rows[0].cells
    add_cell_text(left, "Acta de Vinculación al Piloto · Modelo 1.0", color=GRAY, size=8.5)
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Página ")
    set_run_font(r, size=8.5, color=GRAY)
    add_word_field(p, "PAGE")
    r = p.add_run(" de ")
    set_run_font(r, size=8.5, color=GRAY)
    add_word_field(p, "NUMPAGES")
    set_table_geometry(table, [6500, 2860], indent=0)


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(5)
    if not CORPORATE_LOGO.is_file():
        raise FileNotFoundError(f"No se encontro el logotipo corporativo: {CORPORATE_LOGO}")
    shape = p.add_run().add_picture(str(CORPORATE_LOGO), width=Inches(2.6))
    shape._inline.docPr.set("title", "Roustix")
    shape._inline.docPr.set("descr", "Logotipo corporativo de Roustix")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Acta de Vinculación al Piloto Controlado")
    set_run_font(r, size=26, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("Modelo operativo para incorporación, aceptación y cierre UAT de empresas piloto")
    set_run_font(r, size=12.5, color=GRAY)

    table = doc.add_table(rows=2, cols=4)
    values = [
        ("Código", "PIL-ACT-001", "Versión", "1.0"),
        ("Estado", "Borrador para completar", "Fecha", "____ / ____ / ______"),
    ]
    for row, vals in zip(table.rows, values):
        for idx, value in enumerate(vals):
            add_cell_text(row.cells[idx], value, bold=idx % 2 == 0, color=NAVY if idx % 2 == 0 else BLACK)
            if idx % 2 == 0:
                set_cell_shading(row.cells[idx], LIGHT_BLUE)
    set_table_geometry(table, [1200, 3060, 1200, 3900])
    style_table(table, header=False, label_columns={0, 2})

    doc.add_paragraph()
    add_callout(
        doc,
        "Documento operativo del piloto",
        "Esta acta registra condiciones de vinculación y evidencia UAT. No reemplaza los términos del servicio, la política de datos, el acuerdo de transmisión ni la asesoría jurídica aplicable.",
        fill="FFF8E8",
        accent=GOLD,
    )


def build() -> Path:
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    enable_field_updates(doc)
    add_title_block(doc)

    add_heading(doc, "1. Identificación de la empresa piloto")
    add_label_detail_table(doc, [
        ("Razón social", "________________________________________________________________"),
        ("NIT / identificación", "________________________________________________________________"),
        ("Representante autorizado", "________________________________________________________________"),
        ("Cargo e identificación", "________________________________________________________________"),
        ("Domicilio y dirección", "________________________________________________________________"),
        ("Correo y teléfono", "________________________________________________________________"),
        ("Sector / actividad", "________________________________________________________________"),
    ])

    add_heading(doc, "2. Condiciones de vinculación")
    add_label_detail_table(doc, [
        ("Plan", "☐ Start    ☐ Business    ☐ Enterprise    ☐ Otro: __________________"),
        ("Modalidad", "☐ Gratuito    ☐ Pago manual    ☐ Otro: __________________________"),
        ("Inicio y final previsto", "Desde ____ / ____ / ______ hasta ____ / ____ / ______"),
        ("Usuarios autorizados", "________ usuarios · límite acordado: __________________________"),
        ("Almacenamiento", "Cuota base: __________ GB · add-on: __________ GB"),
        ("Módulos habilitados", "_______________________________________________________________"),
        ("Valor e impuestos", "$ __________________ COP · condición de pago: __________________"),
        ("Tenant / slug", "_______________________________________________________________"),
    ])

    add_heading(doc, "3. Objetivo y alcance del piloto")
    p = doc.add_paragraph(
        "El piloto busca validar que la empresa pueda configurar y operar Roustix con sus propios usuarios y datos, comprobar los flujos esenciales y registrar hallazgos antes de una disponibilidad comercial general."
    )
    p.paragraph_format.keep_together = True
    add_checkbox(doc, "Onboarding asistido y configuración inicial de empresa, sede, áreas y usuarios.")
    add_checkbox(doc, "Registro y gestión de activos, ficha técnica, hoja de vida y documentos.")
    add_checkbox(doc, "Incidencia, asignación, orden de trabajo, jornadas, repuestos e informe técnico.")
    add_checkbox(doc, "Notificaciones, almacenamiento, soporte y exportaciones habilitadas por el plan.")
    add_checkbox(doc, "Otros módulos expresamente acordados: ____________________________________________.")

    add_heading(doc, "4. Responsables del recorrido")
    table = doc.add_table(rows=1, cols=4)
    for cell, value in zip(table.rows[0].cells, ("Rol", "Nombre", "Correo / teléfono", "Responsabilidad")):
        add_cell_text(cell, value, bold=True, color=WHITE)
    roles = [
        ("Empresa · líder", "", "", "Coordina usuarios, datos y aceptación"),
        ("Empresa · administrador", "", "", "Configura tenant y permisos"),
        ("Empresa · técnico", "", "", "Ejecuta el flujo de mantenimiento"),
        ("Roustix · onboarding", "", "", "Guía, registra evidencia y soporte"),
        ("Roustix · escalamiento", "", "", "Atiende bloqueos o incidentes"),
    ]
    for role in roles:
        cells = table.add_row().cells
        for cell, value in zip(cells, role):
            add_cell_text(cell, value)
    set_table_geometry(table, [1800, 2100, 2500, 2960])
    style_table(table)

    doc.add_page_break()
    add_heading(doc, "5. UAT abreviado de la empresa piloto")
    p = doc.add_paragraph(
        "Marque el resultado de cada comprobación y anote una evidencia breve. La infraestructura general ya fue certificada; este recorrido valida la configuración particular de la empresa."
    )
    p.paragraph_format.keep_with_next = True

    table = doc.add_table(rows=1, cols=4)
    headers = ("N.º", "Comprobación", "Resultado", "Evidencia / observación")
    for cell, value in zip(table.rows[0].cells, headers):
        add_cell_text(cell, value, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    checks = [
        "Registro, verificación e inicio de sesión",
        "Empresa, sector, zona horaria, sede y área",
        "Usuarios, cargos, roles y permisos",
        "Activo con responsable, proveedor, factura e imagen",
        "Ficha técnica y campos sectoriales",
        "Hoja de vida y PDF con datos y hora correctos",
        "Incidencia y notificación en campana",
        "Asignación y creación de OT desde incidencia",
        "Jornada, recibido por y paro cuando aplique",
        "Repuesto e informe técnico cuando aplique",
        "Cierre de OT e incidencia asociada",
        "Almacenamiento, plan, límites y soporte",
        "Aislamiento: acceso exclusivo a datos del tenant",
    ]
    for idx, check in enumerate(checks, 1):
        cells = table.add_row().cells
        add_cell_text(cells[0], str(idx), align=WD_ALIGN_PARAGRAPH.CENTER)
        add_cell_text(cells[1], check)
        add_cell_text(cells[2], "☐ OK\n☐ Obs.\n☐ Falla", size=9)
        add_cell_text(cells[3], "")
    set_table_geometry(table, [600, 3900, 1500, 3360])
    style_table(table)

    add_heading(doc, "6. Hallazgos y acciones")
    table = doc.add_table(rows=1, cols=5)
    for cell, value in zip(table.rows[0].cells, ("ID", "Hallazgo", "Severidad", "Responsable", "Fecha / estado")):
        add_cell_text(cell, value, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    for idx in range(1, 6):
        cells = table.add_row().cells
        add_cell_text(cells[0], f"H-{idx:02}", align=WD_ALIGN_PARAGRAPH.CENTER)
        for cell in cells[1:]:
            add_cell_text(cell, "")
    set_table_geometry(table, [700, 3600, 1400, 1800, 1860])
    style_table(table)

    add_callout(
        doc,
        "Criterio de bloqueo",
        "Una pérdida o exposición de datos, falla de aislamiento, imposibilidad de recuperar acceso o bloqueo de un flujo esencial impide aprobar la vinculación hasta su corrección.",
        fill="FFF0F0",
        accent=RED,
    )

    add_heading(doc, "7. Soporte, comunicación y escalamiento")
    add_label_detail_table(doc, [
        ("Canal principal", "_______________________________________________________________"),
        ("Horario", "_______________________________________________________________"),
        ("Primera respuesta objetivo", "Crítica: ______ · Alta: ______ · Normal: ______"),
        ("Contacto del cliente", "_______________________________________________________________"),
        ("Contacto Roustix", "_______________________________________________________________"),
        ("Canal de incidentes críticos", "_______________________________________________________________"),
    ])

    add_heading(doc, "8. Datos, almacenamiento y continuidad")
    add_checkbox(doc, "La empresa declara que está autorizada para incorporar los datos del piloto.")
    add_checkbox(doc, "Los usuarios conocen sus responsabilidades de acceso y confidencialidad.")
    add_checkbox(doc, "La empresa conoce su cuota de almacenamiento y las reglas al 80 % y 100 %.")
    add_checkbox(doc, "Se informó el procedimiento de exportación, respaldo, soporte y reporte de incidentes.")
    add_checkbox(doc, "Se entregaron o aceptaron los documentos de privacidad y tratamiento aplicables.")

    add_heading(doc, "9. Resultado del recorrido")
    add_label_detail_table(doc, [
        ("Decisión", "☐ APROBADO    ☐ APROBADO CON OBSERVACIONES    ☐ NO APROBADO"),
        ("Fecha y versión Roustix", "____ / ____ / ______ · v________________ · commit ________________"),
        ("Inicio comercial / piloto", "☐ Autorizado    ☐ Condicionado    ☐ No autorizado"),
        ("Próxima revisión", "____ / ____ / ______"),
    ])
    add_field_line(doc, "Observación general")
    add_field_line(doc, "Condiciones pendientes")

    add_heading(doc, "10. Aceptación y firmas")
    p = doc.add_paragraph(
        "Las partes confirman que la información consignada refleja las condiciones operativas del piloto y el resultado observado. La firma no sustituye los documentos legales o comerciales que correspondan."
    )
    p.paragraph_format.keep_together = True

    table = doc.add_table(rows=5, cols=2)
    labels = (
        ("POR LA EMPRESA PILOTO", "POR ROUSTIX · OPERADOR 1"),
        ("Firma: __________________________", "Firma: __________________________"),
        ("Nombre: ________________________", "Nombre: ________________________"),
        ("Cargo / ID: _____________________", "Identificación: __________________"),
        ("Fecha: __________________________", "Fecha: __________________________"),
    )
    for row, vals in zip(table.rows, labels):
        for idx, value in enumerate(vals):
            add_cell_text(row.cells[idx], value, bold=row._index == 0, color=NAVY if row._index == 0 else BLACK)
            if row._index == 0:
                set_cell_shading(row.cells[idx], LIGHT_BLUE)
    set_table_geometry(table, [4680, 4680])
    style_table(table, header=False)

    doc.add_paragraph()
    table = doc.add_table(rows=5, cols=2)
    labels = (
        ("POR ROUSTIX · OPERADOR 2", "CONTROL DOCUMENTAL"),
        ("Firma: __________________________", "Elaboró: _________________________"),
        ("Nombre: ________________________", "Revisó: __________________________"),
        ("Identificación: __________________", "Aprobó: __________________________"),
        ("Fecha: __________________________", "Fecha: __________________________"),
    )
    for row, vals in zip(table.rows, labels):
        for idx, value in enumerate(vals):
            add_cell_text(row.cells[idx], value, bold=row._index == 0, color=NAVY if row._index == 0 else BLACK)
            if row._index == 0:
                set_cell_shading(row.cells[idx], LIGHT_BLUE)
    set_table_geometry(table, [4680, 4680])
    style_table(table, header=False)

    add_callout(
        doc,
        "Archivo de evidencia",
        "Conserve esta acta junto con capturas saneadas, reporte de hallazgos y evidencia del UAT. No adjunte contraseñas, códigos, tokens, cookies ni secretos de infraestructura.",
        fill=PALE_BLUE,
        accent=BLUE,
    )

    props = doc.core_properties
    props.title = "Acta de Vinculación al Piloto Controlado de Roustix"
    props.subject = "Modelo operativo para vinculación y UAT abreviado de empresas piloto"
    props.author = "Roustix"
    props.keywords = "Roustix, piloto, UAT, vinculación, onboarding"
    props.comments = "Modelo genérico sin datos personales. Requiere completar y revisar antes de firmar."

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
