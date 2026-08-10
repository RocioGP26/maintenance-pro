"""Genera documentos comerciales de cotización, incorporación y cierre Roustix."""

from __future__ import annotations

import os
from pathlib import Path

from docx import Document

from build_pilot_document_package import (
    GOLD,
    add_callout,
    add_heading,
    add_key_values,
    add_matrix,
    add_para,
    add_signature_table,
    start_document,
)


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = Path(
    os.environ.get(
        "ROUSTIX_COMMERCIAL_OPERATIONS_OUTPUT_DIR",
        ROOT / "docs" / "production-readiness" / "templates",
    )
)


def save_document(doc: Document, filename: str, title: str, subject: str) -> Path:
    props = doc.core_properties
    props.title = title
    props.subject = subject
    props.author = "Roustix"
    props.keywords = "Roustix, SaaS, comercial, onboarding, capacitación, aceptación, cierre"
    props.comments = "Plantilla corporativa para completar; no sustituye revisión jurídica o contable."
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / filename
    doc.save(path)
    return path


def build_quote() -> Path:
    doc = start_document(
        "Cotización Comercial Roustix",
        "Oferta económica individual y condiciones de validez",
        "COM-COT-001",
        "Plantilla comercial para completar",
        "Presentar al Cliente una oferta económica identificable, verificable y separada de la propuesta narrativa, con plan, complementos, impuestos y condiciones de pago.",
        version="1.0",
    )
    add_callout(
        doc,
        "Control de emisión",
        "No emitir como definitiva hasta completar la identidad del Prestador, el tratamiento tributario, la vigencia, el medio de pago y la aprobación comercial interna.",
        fill="FFF8E8",
        accent=GOLD,
    )
    add_heading(doc, "1. Identificación")
    add_key_values(doc, [
        ("Cotización No.", "RTX-COT-____________"),
        ("Fecha de emisión", "____ / ____ / ______"),
        ("Válida hasta", "____ / ____ / ______"),
        ("Cliente / razón social", "____________________________________________________________"),
        ("NIT / identificación", "____________________________________________________________"),
        ("Contacto / cargo", "____________________________________________________________"),
        ("Correo / teléfono", "____________________________________________________________"),
        ("Propuesta relacionada", "Código / versión: ___________________________________________"),
    ])
    add_heading(doc, "2. Suscripción y complementos")
    add_matrix(doc, ("Ítem", "SKU", "Cant.", "Periodicidad", "Valor unitario", "Subtotal"), [
        ("Plan Roustix: __________________", "PLAN-________", "1", "Mensual", "$ __________", "$ __________"),
        ("Almacenamiento adicional", "ADD-STG-____", "____", "Mensual", "$ __________", "$ __________"),
        ("Usuarios adicionales", "ADD-USR-COT", "____", "Mensual", "$ __________", "$ __________"),
        ("Soporte / SLA adicional", "ADD-SUP-COT", "____", "Mensual", "$ __________", "$ __________"),
        ("Otro: ________________________", "____________", "____", "________", "$ __________", "$ __________"),
    ], [2600, 1500, 800, 1500, 1480, 1480], size=8.8)
    add_heading(doc, "3. Servicios profesionales")
    add_matrix(doc, ("Servicio", "Alcance / unidad", "Cant.", "Valor unitario", "Subtotal"), [
        ("Onboarding / implementación", "____________________________", "____", "$ __________", "$ __________"),
        ("Migración o depuración de datos", "____________________________", "____", "$ __________", "$ __________"),
        ("Capacitación adicional", "____________________________", "____", "$ __________", "$ __________"),
        ("Integración / desarrollo especial", "____________________________", "____", "$ __________", "$ __________"),
        ("Otro: ________________________", "____________________________", "____", "$ __________", "$ __________"),
    ], [2600, 2900, 800, 1530, 1530], size=8.8)
    add_heading(doc, "4. Resumen económico")
    add_key_values(doc, [
        ("Subtotal recurrente", "$ __________________________ COP"),
        ("Subtotal pago único", "$ __________________________ COP"),
        ("Descuento autorizado", "$ __________________________ COP / ______ %"),
        ("Impuestos y retenciones", "____________________________________________________________"),
        ("Total inicial", "$ __________________________ COP"),
        ("Total periódico", "$ __________________________ COP"),
        ("Periodicidad", "☐ Mensual anticipada  ☐ Trimestral  ☐ Anual  ☐ Otra: ______________"),
        ("Medio de pago", "____________________________________________________________"),
        ("Vencimiento", "________ días desde emisión / fecha: ____ / ____ / ______"),
    ])
    add_heading(doc, "5. Alcance y exclusiones")
    add_matrix(doc, ("Concepto", "Definición acordada"), [
        ("Incluye", "Plan, módulos, capacidad y servicios descritos en esta cotización y su propuesta relacionada."),
        ("No incluye", "Desarrollos, desplazamientos, migraciones, impuestos o servicios no indicados expresamente."),
        ("Dependencias del Cliente", "Entrega oportuna de datos, responsables, accesos autorizados y validaciones requeridas."),
        ("Cambios de alcance", "Requieren cotización u Orden adicional aceptada por ambas partes."),
    ], [2600, 6760])
    add_heading(doc, "6. Condiciones de aceptación")
    add_para(doc, "La aceptación de esta cotización confirma únicamente los valores y el alcance económico aquí descritos. La prestación se rige además por la propuesta aceptada, el Contrato SaaS, el Acta de Servicio, los documentos de privacidad, la Política de Soporte y el SLA cuando se anexe expresamente.")
    add_key_values(doc, [
        ("Orden / acta que se emitirá", "____________________________________________________________"),
        ("Fecha estimada de activación", "____ / ____ / ______"),
        ("Aprobación interna Roustix", "Nombre / cargo / fecha: __________________________________________"),
    ])
    add_heading(doc, "7. Aceptación")
    add_signature_table(doc, "Empresa Cliente", "Prestador Roustix")
    return save_document(doc, "12-COTIZACION-COMERCIAL-ROUSTIX.docx", "Cotización Comercial Roustix", "Oferta económica individual del servicio SaaS")


def build_onboarding_guide() -> Path:
    doc = start_document(
        "Guía de Onboarding Comercial",
        "Preparación, configuración, capacitación y activación del servicio",
        "OPS-ONB-001",
        "Guía operativa vigente",
        "Guiar al Cliente y a Roustix desde la aceptación comercial hasta la activación controlada, dejando responsables, datos, seguridad y criterios de salida claramente documentados.",
        version="1.0",
    )
    add_heading(doc, "1. Datos de la incorporación")
    add_key_values(doc, [
        ("Empresa / tenant", "____________________________________________________________"),
        ("Plan y módulos", "____________________________________________________________"),
        ("Líder del Cliente", "____________________________________________________________"),
        ("Líder Roustix", "____________________________________________________________"),
        ("Inicio previsto", "____ / ____ / ______"),
        ("Activación prevista", "____ / ____ / ______"),
        ("Modalidad", "☐ Remota  ☐ Presencial cotizada  ☐ Híbrida"),
    ])
    add_heading(doc, "2. Responsabilidades")
    add_matrix(doc, ("Rol", "Responsabilidad mínima", "Nombre / contacto"), [
        ("Patrocinador del Cliente", "Aprueba alcance, responsables y decisiones de negocio.", ""),
        ("Administrador", "Usuarios, roles, sedes, áreas y configuración autorizada.", ""),
        ("Supervisor", "Valida procesos, datos iniciales y criterios operativos.", ""),
        ("Técnico / usuario clave", "Ejecuta recorridos funcionales y reporta hallazgos.", ""),
        ("Onboarding Roustix", "Coordina configuración, capacitación y gate de activación.", ""),
    ], [2100, 4300, 2960])
    add_heading(doc, "3. Ruta de incorporación")
    add_matrix(doc, ("Etapa", "Acciones", "Entregable", "Estado"), [
        ("1. Preparación", "Confirmar documentos, plan, responsables y cronograma.", "Expediente comercial completo", "☐"),
        ("2. Datos", "Preparar sedes, áreas, usuarios, activos y catálogos.", "Plantillas validadas", "☐"),
        ("3. Configuración", "Crear tenant, módulos, roles, cuotas y parámetros.", "Entorno configurado", "☐"),
        ("4. Capacitación", "Formar perfiles y practicar los flujos acordados.", "Acta de capacitación", "☐"),
        ("5. Aceptación", "Ejecutar controles y registrar pendientes.", "Acta de entrega y aceptación", "☐"),
        ("6. Activación", "Habilitar operación y comunicar soporte.", "Servicio activo", "☐"),
        ("7. Seguimiento", "Revisar adopción, incidencias y ajustes autorizados.", "Minuta de seguimiento", "☐"),
    ], [1600, 3600, 2860, 1300], size=8.8)
    add_heading(doc, "4. Información que debe preparar el Cliente")
    add_matrix(doc, ("Conjunto", "Contenido", "Responsable", "Validado"), [
        ("Organización", "Sedes, áreas, ubicaciones, procesos y centros de costo.", "", "☐"),
        ("Usuarios", "Nombre, correo, cargo, área, rol y estado.", "", "☐"),
        ("Activos", "Código, nombre, tipo, ubicación, responsable y datos técnicos.", "", "☐"),
        ("Mantenimiento", "Planes, frecuencias, actividades, técnicos y repuestos.", "", "☐"),
        ("Archivos", "Logos, fotografías, fichas, informes y documentos autorizados.", "", "☐"),
    ], [1900, 4400, 1900, 1160], size=8.8)
    add_callout(doc, "Protección de información", "No enviar contraseñas, tokens, claves privadas ni datos sensibles innecesarios. Los archivos deben contar con autorización y respetar la cuota contratada.", fill="FFF8E8", accent=GOLD)
    add_heading(doc, "5. Configuración y seguridad")
    add_matrix(doc, ("Control", "Criterio de salida", "Estado / evidencia"), [
        ("Tenant y plan", "Empresa, plan, módulos y cuota corresponden al Acta de Servicio.", ""),
        ("Administrador", "Cuenta verificada, MFA si aplica y datos de recuperación válidos.", ""),
        ("Roles", "Mínimo privilegio; administradores, supervisores, técnicos y reportantes diferenciados.", ""),
        ("Sesiones", "Acceso probado y procedimiento de recuperación conocido.", ""),
        ("Privacidad", "Documentos entregados y aceptación registrada cuando corresponda.", ""),
        ("Soporte", "Canal, horario, severidades y escalamiento comunicados.", ""),
    ], [2200, 4700, 2460])
    add_heading(doc, "6. Capacitación mínima por perfil")
    add_matrix(doc, ("Perfil", "Debe poder realizar"), [
        ("Administrador", "Gestionar usuarios, roles, empresa, sedes, áreas, catálogos y almacenamiento."),
        ("Supervisor", "Revisar incidencias, asignaciones, OT, jornadas, repuestos, costos e indicadores."),
        ("Técnico", "Consultar trabajo asignado, registrar avances, evidencias, repuestos y cierres autorizados."),
        ("Reportante", "Crear incidencias, adjuntar información, consultar estado y responder solicitudes."),
    ], [2200, 7160])
    add_heading(doc, "7. Gate de activación")
    add_matrix(doc, ("Verificación", "Resultado", "Observación / evidencia"), [
        ("Documentos comerciales y de datos completos", "☐ Aprobado  ☐ No", ""),
        ("Pago o condición comercial validada", "☐ Aprobado  ☐ No", ""),
        ("Configuración y datos iniciales validados", "☐ Aprobado  ☐ No", ""),
        ("Capacitación realizada", "☐ Aprobado  ☐ No", ""),
        ("Recorridos funcionales aceptados", "☐ Aprobado  ☐ No", ""),
        ("Soporte y escalamiento comunicados", "☐ Aprobado  ☐ No", ""),
        ("Capacidad operativa Roustix disponible", "☐ Aprobado  ☐ No", ""),
    ], [4300, 2100, 2960])
    add_heading(doc, "8. Seguimiento inicial")
    add_matrix(doc, ("Momento", "Objetivo", "Fecha / responsable"), [
        ("Día 2-3", "Confirmar accesos y resolver bloqueos iniciales.", ""),
        ("Día 7", "Revisar adopción, calidad de datos e incidencias.", ""),
        ("Día 15", "Validar flujos principales y ajustes autorizados.", ""),
        ("Día 30", "Revisar uso, resultados, soporte y siguiente plan de trabajo.", ""),
    ], [1700, 4700, 2960])
    add_heading(doc, "9. Canales")
    add_key_values(doc, [
        ("Correo corporativo", "contacto@roustix.com"),
        ("Horario base", "Lunes a viernes, 8:00 a. m. a 5:00 p. m. - America/Bogota"),
        ("Canal especial contratado", "____________________________________________________________"),
    ])
    return save_document(doc, "13-GUIA-ONBOARDING-COMERCIAL-ROUSTIX.docx", "Guía de Onboarding Comercial Roustix", "Preparación y activación del servicio")


def build_training_act() -> Path:
    doc = start_document(
        "Acta de Capacitación Roustix",
        "Asistencia, contenidos, prácticas y compromisos",
        "OPS-CAP-001",
        "Plantilla operativa para completar",
        "Dejar evidencia de la capacitación impartida, los perfiles participantes, los contenidos tratados, las prácticas realizadas y los compromisos posteriores.",
        version="1.0",
    )
    add_heading(doc, "1. Identificación de la sesión")
    add_key_values(doc, [
        ("Empresa / tenant", "____________________________________________________________"),
        ("Fecha", "____ / ____ / ______"),
        ("Hora", "Inicio ______  Fin ______  Zona __________________"),
        ("Modalidad / lugar", "☐ Remota  ☐ Presencial  ☐ Híbrida: ________________________"),
        ("Facilitador Roustix", "____________________________________________________________"),
        ("Objetivo", "____________________________________________________________"),
        ("Módulos cubiertos", "____________________________________________________________"),
    ])
    add_heading(doc, "2. Participantes")
    add_matrix(doc, ("No.", "Nombre", "Cargo / área", "Perfil Roustix", "Correo", "Firma / asistencia"), [
        (str(i), "", "", "", "", "") for i in range(1, 9)
    ], [600, 1700, 1700, 1600, 2200, 1560], size=8.4)
    add_heading(doc, "3. Contenidos impartidos")
    add_matrix(doc, ("Bloque", "Contenido", "Perfil objetivo", "Cubierto"), [
        ("Acceso y seguridad", "Ingreso, recuperación, sesiones, confidencialidad y responsabilidades.", "Todos", "☐"),
        ("Administración", "Usuarios, roles, empresa, sedes, áreas, catálogos y almacenamiento.", "Administrador", "☐"),
        ("Supervisión", "Incidencias, asignación, OT, jornadas, repuestos, costos e indicadores.", "Supervisor", "☐"),
        ("Ejecución", "Trabajo asignado, avances, evidencias, repuestos y cierre.", "Técnico", "☐"),
        ("Reporte", "Creación de incidencias, adjuntos, seguimiento y respuesta.", "Reportante", "☐"),
        ("Soporte", "Canales, información mínima, severidades y escalamiento.", "Todos", "☐"),
    ], [1700, 4600, 1800, 1260], size=8.7)
    add_heading(doc, "4. Prácticas y verificación")
    add_matrix(doc, ("Práctica", "Resultado esperado", "Resultado", "Observación"), [
        ("Ingreso y perfil", "El participante accede y reconoce sus permisos.", "☐ Cumple  ☐ No", ""),
        ("Flujo principal", "Ejecuta el recorrido correspondiente a su rol.", "☐ Cumple  ☐ No", ""),
        ("Archivos y datos", "Carga información autorizada y conoce límites.", "☐ Cumple  ☐ No", ""),
        ("Soporte", "Sabe reportar un incidente sin compartir secretos.", "☐ Cumple  ☐ No", ""),
    ], [2300, 3800, 1800, 1460], size=8.7)
    add_heading(doc, "5. Material entregado")
    add_key_values(doc, [
        ("Guía / manual", "____________________________________________________________"),
        ("Grabación autorizada", "☐ No  ☐ Sí, ubicación / vigencia: ______________________________"),
        ("Ejercicios / archivos", "____________________________________________________________"),
        ("Política de soporte", "☐ Entregada  ☐ Referenciada  ☐ Pendiente"),
    ])
    add_heading(doc, "6. Preguntas, hallazgos y compromisos")
    add_matrix(doc, ("No.", "Pregunta / compromiso", "Responsable", "Fecha", "Estado"), [
        (str(i), "", "", "", "") for i in range(1, 6)
    ], [650, 3800, 1900, 1600, 1410])
    add_heading(doc, "7. Evaluación de la sesión")
    add_key_values(doc, [
        ("Comprensión general", "☐ Alta  ☐ Adecuada  ☐ Requiere refuerzo"),
        ("Próxima sesión / refuerzo", "☐ No requerido  ☐ Sí: ____ / ____ / ______"),
        ("Observaciones", "____________________________________________________________"),
    ])
    add_heading(doc, "8. Constancia")
    add_para(doc, "Las firmas confirman la realización de la sesión y la asistencia registrada. No sustituyen la aceptación contractual ni garantizan dominio total de funciones no practicadas.")
    add_signature_table(doc, "Representante del Cliente", "Facilitador Roustix")
    return save_document(doc, "14-ACTA-CAPACITACION-ROUSTIX.docx", "Acta de Capacitación Roustix", "Asistencia y resultados de capacitación")


def build_acceptance_act() -> Path:
    doc = start_document(
        "Acta de Entrega y Aceptación",
        "Configuración, recorridos funcionales y autorización de activación",
        "OPS-ACE-001",
        "Plantilla operativa para completar y firmar",
        "Registrar qué fue entregado, qué fue probado, cuáles pendientes permanecen y si el Cliente autoriza la activación del servicio contratado.",
        version="1.0",
    )
    add_heading(doc, "1. Identificación")
    add_key_values(doc, [
        ("Empresa / tenant", "____________________________________________________________"),
        ("Contrato / Acta de Servicio", "____________________________________________________________"),
        ("Plan y módulos", "____________________________________________________________"),
        ("Responsable del Cliente", "____________________________________________________________"),
        ("Responsable Roustix", "____________________________________________________________"),
        ("Fecha de revisión", "____ / ____ / ______"),
    ])
    add_heading(doc, "2. Entregables revisados")
    add_matrix(doc, ("Entregable", "Alcance verificado", "Resultado", "Evidencia / observación"), [
        ("Tenant y plan", "Empresa, plan, módulos y capacidad contratada.", "☐ Conforme  ☐ No", ""),
        ("Organización", "Sedes, áreas, ubicaciones y responsables iniciales.", "☐ Conforme  ☐ No", ""),
        ("Usuarios y roles", "Cuentas activas y permisos por perfil.", "☐ Conforme  ☐ No", ""),
        ("Datos iniciales", "Activos, catálogos y documentos incluidos en el alcance.", "☐ Conforme  ☐ No", ""),
        ("Flujos", "Incidencias, OT, jornadas, repuestos y reportes acordados.", "☐ Conforme  ☐ No", ""),
        ("Capacitación", "Sesiones y material indicados en el acta relacionada.", "☐ Conforme  ☐ No", ""),
        ("Soporte", "Canales, horarios y escalamiento comunicados.", "☐ Conforme  ☐ No", ""),
    ], [2200, 3600, 1900, 1660], size=8.5)
    add_heading(doc, "3. Recorridos de aceptación")
    add_matrix(doc, ("No.", "Recorrido", "Perfil", "Resultado", "Evidencia"), [
        ("1", "Acceso, recuperación y permisos", "Administrador", "☐ Aprobado  ☐ No", ""),
        ("2", "Configuración y administración básica", "Administrador", "☐ Aprobado  ☐ No", ""),
        ("3", "Flujo operativo principal", "Supervisor / técnico", "☐ Aprobado  ☐ No", ""),
        ("4", "Reporte y seguimiento", "Reportante", "☐ Aprobado  ☐ No", ""),
        ("5", "Exportación / documento clave", "Responsable", "☐ Aprobado  ☐ No", ""),
    ], [650, 3400, 2000, 1900, 1410], size=8.5)
    add_heading(doc, "4. Pendientes aceptados")
    add_matrix(doc, ("ID", "Pendiente", "Severidad", "Responsable", "Fecha objetivo", "Bloquea"), [
        ("", "", "", "", "", "☐ Sí  ☐ No") for _ in range(5)
    ], [650, 3000, 1200, 1700, 1500, 1310], size=8.5)
    add_heading(doc, "5. Decisión")
    add_key_values(doc, [
        ("Resultado", "☐ Aceptación sin reservas  ☐ Aceptación con pendientes no bloqueantes  ☐ No aceptación"),
        ("Activación autorizada", "☐ Sí  ☐ No"),
        ("Fecha efectiva", "____ / ____ / ______  Hora ______  Zona __________________"),
        ("Próxima revisión", "____ / ____ / ______"),
        ("Motivo si no se acepta", "____________________________________________________________"),
    ])
    add_callout(doc, "Alcance de la firma", "La aceptación confirma los entregables y recorridos descritos. No elimina garantías, obligaciones de soporte ni pendientes expresamente registrados.", fill="FFF8E8", accent=GOLD)
    add_heading(doc, "6. Firmas")
    add_signature_table(doc, "Empresa Cliente", "Prestador Roustix")
    return save_document(doc, "15-ACTA-ENTREGA-ACEPTACION-ROUSTIX.docx", "Acta de Entrega y Aceptación Roustix", "Entrega y autorización de activación")


def build_termination_act() -> Path:
    doc = start_document(
        "Acta de Terminación del Servicio",
        "Cierre contractual, accesos, datos y obligaciones posteriores",
        "OPS-CIE-001",
        "Plantilla comercial para completar y firmar",
        "Formalizar el cierre total o parcial del servicio Roustix y evitar que permanezcan accesos, saldos, datos, integraciones o compromisos sin una decisión documentada.",
        version="1.0",
    )
    add_heading(doc, "1. Identificación")
    add_key_values(doc, [
        ("Empresa / tenant", "____________________________________________________________"),
        ("Contrato / Orden", "____________________________________________________________"),
        ("Solicitud de terminación", "____________________________________________________________"),
        ("Tipo de cierre", "☐ Total  ☐ Parcial  ☐ No renovación  ☐ Terminación anticipada"),
        ("Fecha efectiva", "____ / ____ / ______  Hora ______  Zona __________________"),
        ("Motivo", "____________________________________________________________"),
    ])
    add_heading(doc, "2. Cierre comercial y financiero")
    add_matrix(doc, ("Concepto", "Estado", "Valor / fecha / evidencia"), [
        ("Facturas emitidas", "☐ Al día  ☐ Pendiente", ""),
        ("Saldo final", "☐ $0  ☐ Pendiente", ""),
        ("Créditos / ajustes", "☐ No aplica  ☐ Aplicado", ""),
        ("Renovación", "☐ Cancelada  ☐ No aplica", ""),
        ("Complementos", "☐ Retirados  ☐ Continúan parcialmente", ""),
    ], [2500, 2300, 4560])
    add_heading(doc, "3. Exportación y entrega")
    add_matrix(doc, ("Conjunto", "Formato / medio", "Fecha", "Responsable", "Evidencia"), [
        ("Datos operativos", "", "", "", ""),
        ("Archivos / evidencias", "", "", "", ""),
        ("Reportes / auditoría autorizada", "", "", "", ""),
        ("Otros", "", "", "", ""),
    ], [2100, 2200, 1500, 1700, 1860])
    add_para(doc, "La entrega detallada debe respaldarse con la Constancia de Exportación, Conservación o Eliminación vigente.")
    add_heading(doc, "4. Accesos e integraciones")
    add_matrix(doc, ("Acción", "Responsable", "Fecha", "Estado / evidencia"), [
        ("Bloquear nuevos accesos", "", "", ""),
        ("Revocar usuarios y sesiones", "", "", ""),
        ("Revocar API keys y webhooks", "", "", ""),
        ("Retirar credenciales o integraciones del Cliente", "", "", ""),
        ("Cerrar canales dedicados", "", "", ""),
    ], [3500, 1800, 1500, 2560])
    add_heading(doc, "5. Conservación y eliminación")
    add_matrix(doc, ("Categoría", "Decisión", "Plazo / fecha", "Fundamento / evidencia"), [
        ("Base activa", "☐ Exportar  ☐ Eliminar  ☐ Conservar", "", ""),
        ("Archivos activos", "☐ Exportar  ☐ Eliminar  ☐ Conservar", "", ""),
        ("Respaldos", "Rotación residual", "", ""),
        ("Auditoría / seguridad", "☐ Conservar  ☐ Anonimizar", "", ""),
        ("Facturación / legal", "☐ Conservar según obligación", "", ""),
    ], [1900, 2700, 1800, 2960])
    add_callout(doc, "Precisión obligatoria", "No declarar eliminación inmediata si existen respaldos en rotación, obligaciones legales o evidencia pendiente. Registrar fechas reales y emitir la constancia de datos correspondiente.", fill="FFF8E8", accent=GOLD)
    add_heading(doc, "6. Pendientes posteriores")
    add_matrix(doc, ("ID", "Compromiso", "Responsable", "Fecha", "Criterio de cierre"), [
        ("", "", "", "", "") for _ in range(5)
    ], [650, 3300, 1800, 1500, 2110])
    add_heading(doc, "7. Obligaciones que sobreviven")
    add_para(doc, "La confidencialidad, propiedad intelectual, tratamiento de datos, pagos, responsabilidades y demás obligaciones que por su naturaleza sobrevivan continuarán conforme al Contrato y la ley aplicable.")
    add_heading(doc, "8. Declaración y firmas")
    add_para(doc, "Las partes dejan constancia de la terminación en la fecha indicada y de las acciones, reservas y pendientes consignados en esta acta.")
    add_signature_table(doc, "Empresa Cliente", "Prestador Roustix")
    return save_document(doc, "16-ACTA-TERMINACION-SERVICIO-ROUSTIX.docx", "Acta de Terminación del Servicio Roustix", "Cierre contractual y tratamiento de datos")


def build_partner_decision_matrix() -> Path:
    doc = start_document(
        "Matriz de Decisiones de Socios",
        "Definiciones jurídicas, comerciales, tributarias y operativas",
        "GOB-DEC-001",
        "Documento interno controlado",
        "Centralizar las decisiones que deben aprobar los socios antes de emitir contratos, precios, SLA o comunicaciones comerciales definitivas.",
        version="1.0",
    )
    add_callout(doc, "Uso interno", "Cada decisión debe registrar responsable, fecha y evidencia. Los asuntos jurídicos o tributarios solo se cierran después del concepto profesional aplicable.", fill="FFF8E8", accent=GOLD)
    add_heading(doc, "1. Control de la sesión")
    add_key_values(doc, [
        ("Fecha / sesión", "____ / ____ / ______  No. ______"),
        ("Socios participantes", "____________________________________________________________"),
        ("Invitados / asesores", "____________________________________________________________"),
        ("Estado del documento", "☐ En preparación  ☐ En revisión  ☐ Aprobado"),
    ])
    add_heading(doc, "2. Criterios de estado")
    add_matrix(doc, ("Estado", "Uso"), [
        ("Propuesta", "Alternativa formulada, todavía sin aprobación."),
        ("Aprobada", "Decisión adoptada por los socios y soportada con evidencia."),
        ("Requiere asesor", "Pendiente de validación jurídica, contable o tributaria."),
        ("Diferida", "No bloquea la comercialización actual; debe tener fecha de revisión."),
    ], [2100, 7260])
    sections = [
        ("3. Identidad y gobierno", [
            ("GOB-01", "Identidad temporal del prestador hasta constituir la sociedad"),
            ("GOB-02", "Tipo societario, razón social, domicilio y representación legal"),
            ("GOB-03", "Participación, aportes, funciones y facultades de cada socio"),
            ("GOB-04", "Titularidad o cesión de marca, software, dominio y demás activos"),
            ("GOB-05", "Cuenta bancaria, autorización para contratar y manejo de fondos"),
        ]),
        ("4. Oferta, precios e impuestos", [
            ("COM-01", "Precios definitivos de Start, Business y Enterprise"),
            ("COM-02", "Servicios adicionales, implementación y almacenamiento"),
            ("COM-03", "Descuentos máximos, aprobaciones y vigencia de cotizaciones"),
            ("COM-04", "Tratamiento de IVA, ICA, retenciones y precios con/sin impuestos"),
            ("COM-05", "Periodicidad, reajuste y moneda de facturación"),
        ]),
        ("5. Cobro y ciclo contractual", [
            ("FIN-01", "Documento de cobro o factura aplicable y responsable de emisión"),
            ("FIN-02", "Pago anticipado, plazo, mora, suspensión y reactivación"),
            ("FIN-03", "Métodos de pago, devoluciones, notas crédito y conciliación"),
            ("CON-01", "Vigencia inicial, renovación, preaviso y terminación"),
            ("CON-02", "Exportación, conservación, eliminación y rotación de respaldos"),
        ]),
        ("6. SLA, soporte y datos", [
            ("OPS-01", "Horario, canales, severidades y responsables de escalamiento"),
            ("OPS-02", "Disponibilidad comprometida, medición, exclusiones y créditos"),
            ("DAT-01", "Roles de tratamiento, canales de privacidad y retención"),
            ("DAT-02", "Subencargados, ubicación, transferencias y datos sensibles"),
            ("DAT-03", "Topes de responsabilidad, confidencialidad y seguridad"),
        ]),
    ]
    for heading, rows in sections:
        add_heading(doc, heading)
        add_matrix(doc, ("ID", "Decisión requerida", "Estado", "Responsable", "Fecha / evidencia"), [
            (item_id, decision, "", "", "") for item_id, decision in rows
        ], [1050, 4000, 1300, 1500, 1510], size=8.5)
    add_heading(doc, "7. Registro de decisiones aprobadas")
    add_matrix(doc, ("ID", "Decisión aprobada", "Asesor consultado", "Fecha", "Evidencia"), [
        ("", "", "", "", "") for _ in range(6)
    ], [900, 3600, 1800, 1400, 1660])
    add_heading(doc, "8. Aprobación")
    add_para(doc, "La firma confirma únicamente las decisiones identificadas como aprobadas y no sustituye los conceptos profesionales expresamente requeridos.")
    add_signature_table(doc, "Socio 1", "Socio 2")
    return save_document(doc, "17-MATRIZ-DECISIONES-SOCIOS-ROUSTIX.docx", "Matriz de Decisiones de Socios Roustix", "Gobierno y decisiones previas a comercialización")


def build_review_pack() -> Path:
    doc = start_document(
        "Paquete de Revisión Jurídica y Contable",
        "Inventario, preguntas y constancia de concepto profesional",
        "GOB-REV-001",
        "Borrador interno para revisión profesional",
        "Entregar a los asesores un expediente ordenado y dejar trazabilidad de observaciones, decisiones y versiones aprobadas antes de publicar o firmar.",
        version="1.0",
    )
    add_callout(doc, "Advertencia", "Este paquete organiza la revisión, pero no constituye asesoría jurídica, contable ni tributaria. La aprobación debe provenir de profesionales habilitados.", fill="FFF8E8", accent=GOLD)
    add_heading(doc, "1. Datos de la revisión")
    add_key_values(doc, [
        ("Prestador / proyecto", "Roustix / _________________________________________________"),
        ("Revisor jurídico", "____________________________________________________________"),
        ("Revisor contable / tributario", "____________________________________________________________"),
        ("Fecha de entrega", "____ / ____ / ______"),
        ("Fecha objetivo", "____ / ____ / ______"),
        ("Carpeta / expediente", "____________________________________________________________"),
    ])
    add_heading(doc, "2. Inventario documental")
    add_matrix(doc, ("Documento", "Código / versión", "Jurídico", "Contable", "Observación"), [
        ("Términos y condiciones", "", "☐", "☐", ""),
        ("Contrato SaaS", "", "☐", "☐", ""),
        ("Privacidad y tratamiento", "", "☐", "☐", ""),
        ("Acuerdo de transmisión", "", "☐", "☐", ""),
        ("Propuesta y cotización", "", "☐", "☐", ""),
        ("Acta de Servicio", "", "☐", "☐", ""),
        ("Política de soporte y SLA", "", "☐", "☐", ""),
        ("Onboarding, aceptación y cierre", "", "☐", "☐", ""),
    ], [2500, 1900, 1100, 1100, 2760], size=8.5)
    add_heading(doc, "3. Preguntas para revisión jurídica")
    add_matrix(doc, ("Tema", "Pregunta / validación", "Concepto / ajuste requerido"), [
        ("Prestador", "¿La identidad y capacidad para contratar están correctamente expresadas?", ""),
        ("Contratación", "¿Aceptación, firma, evidencia y versiones producen trazabilidad suficiente?", ""),
        ("Responsabilidad", "¿Topes, exclusiones, garantías e indemnidad son válidos y equilibrados?", ""),
        ("Propiedad intelectual", "¿Marca, software, contenidos y licencias tienen titularidad documentada?", ""),
        ("Datos personales", "¿Roles, autorizaciones, transmisión, transferencias y derechos están cubiertos?", ""),
        ("Terminación", "¿Preavisos, suspensión, exportación, retención y eliminación son coherentes?", ""),
        ("Ley aplicable", "¿Jurisdicción, solución de controversias y notificaciones son apropiadas?", ""),
    ], [1900, 4100, 3360], size=8.5)
    add_heading(doc, "4. Preguntas para revisión contable y tributaria")
    add_matrix(doc, ("Tema", "Pregunta / validación", "Concepto / ajuste requerido"), [
        ("Identificación", "¿RUT, responsabilidades, actividad económica y obligaciones están definidas?", ""),
        ("Facturación", "¿Cuándo procede factura electrónica, cuenta de cobro u otro soporte?", ""),
        ("Impuestos", "¿IVA, ICA y retenciones aplican a suscripción, implementación y complementos?", ""),
        ("Precios", "¿Las ofertas deben expresar valores antes o después de impuestos?", ""),
        ("Ingresos", "¿Periodicidad, anticipos, renovaciones y reconocimiento están correctamente tratados?", ""),
        ("Recaudo", "¿Cuenta, conciliación, mora, devoluciones y notas crédito tienen control suficiente?", ""),
        ("Conservación", "¿Qué documentos y soportes deben conservarse y durante cuánto tiempo?", ""),
    ], [1900, 4100, 3360], size=8.5)
    add_heading(doc, "5. Hallazgos y control de cambios")
    add_matrix(doc, ("ID", "Documento / cláusula", "Hallazgo", "Responsable", "Estado"), [
        ("", "", "", "", "") for _ in range(8)
    ], [650, 2300, 3300, 1700, 1410], size=8.5)
    add_heading(doc, "6. Dictamen de salida")
    add_key_values(doc, [
        ("Resultado jurídico", "☐ Aprobado  ☐ Aprobado con ajustes  ☐ No aprobado"),
        ("Resultado contable", "☐ Aprobado  ☐ Aprobado con ajustes  ☐ No aprobado"),
        ("Versiones autorizadas", "____________________________________________________________"),
        ("Restricciones / reservas", "____________________________________________________________"),
        ("Fecha de próxima revisión", "____ / ____ / ______"),
    ])
    add_heading(doc, "7. Constancias")
    add_signature_table(doc, "Revisor profesional", "Socios Roustix")
    return save_document(doc, "18-PAQUETE-REVISION-JURIDICA-CONTABLE-ROUSTIX.docx", "Paquete de Revisión Jurídica y Contable Roustix", "Revisión profesional del paquete comercial")


def build_contracting_checklist() -> Path:
    doc = start_document(
        "Checklist Maestro de Contratación y Activación",
        "Expediente comercial, gates y autorización de servicio",
        "OPS-CON-001",
        "Lista de control operativa",
        "Asegurar que ninguna empresa sea activada sin identidad, oferta, aceptación, pago, tratamiento de datos, configuración y responsables verificables.",
        version="1.0",
    )
    add_heading(doc, "1. Identificación del expediente")
    add_key_values(doc, [
        ("Cliente / empresa", "____________________________________________________________"),
        ("NIT / identificación", "____________________________________________________________"),
        ("Oportunidad / expediente", "RTX-CLI-_______________________________________________"),
        ("Plan / módulos", "____________________________________________________________"),
        ("Responsable comercial", "____________________________________________________________"),
        ("Responsable de activación", "____________________________________________________________"),
        ("Fecha objetivo", "____ / ____ / ______"),
    ])
    stages = [
        ("2. Calificación y oferta", [
            "Necesidad, alcance, usuarios, activos y capacidad identificados",
            "Propuesta y cotización emitidas con vigencia y versión",
            "Precio, impuestos, complementos y descuentos aprobados internamente",
        ]),
        ("3. Contratación y datos", [
            "Identidad y capacidad del cliente verificadas",
            "Contrato / términos y Acta de Servicio aceptados",
            "Privacidad y acuerdo de transmisión aplicables aceptados",
            "Versiones, fecha, firmantes y evidencia de aceptación archivados",
        ]),
        ("4. Cobro y preparación", [
            "Documento de cobro emitido y condición de pago validada",
            "Pago o autorización comercial excepcional registrada",
            "Administrador, dominio, zona horaria y canales confirmados",
            "Capacidad operativa, almacenamiento y soporte disponibles",
        ]),
        ("5. Configuración y capacitación", [
            "Tenant, plan, módulos, cuotas y seguridad configurados",
            "Usuarios, roles y datos iniciales validados",
            "Capacitación ejecutada y acta archivada",
            "Recorridos funcionales y pendientes registrados",
        ]),
        ("6. Activación y seguimiento", [
            "Acta de Entrega y Aceptación aprobada",
            "Canales, horarios y escalamiento comunicados",
            "Activación autorizada por Cliente y Roustix",
            "Seguimientos de días 2-3, 7, 15 y 30 programados",
        ]),
    ]
    for heading, items in stages:
        add_heading(doc, heading)
        add_matrix(doc, ("Control", "Responsable", "Fecha", "Estado", "Evidencia"), [
            (item, "", "", "☐", "") for item in items
        ], [4000, 1500, 1300, 900, 1660], size=8.5)
    add_heading(doc, "7. Gates de no activación")
    add_matrix(doc, ("Gate obligatorio", "Criterio", "Resultado"), [
        ("Aceptación contractual", "Documento vigente, aceptación y firmantes verificables.", "☐ Cumple  ☐ Bloquea"),
        ("Tratamiento de datos", "Roles, documentos y autorizaciones aplicables completos.", "☐ Cumple  ☐ Bloquea"),
        ("Pago", "Condición comercial aprobada y soporte disponible.", "☐ Cumple  ☐ Bloquea"),
        ("Administrador", "Cuenta válida, verificada y con recuperación disponible.", "☐ Cumple  ☐ Bloquea"),
        ("Operación", "Capacidad, configuración, soporte y responsables disponibles.", "☐ Cumple  ☐ Bloquea"),
    ], [2400, 4300, 2660])
    add_callout(doc, "Regla de salida", "Si un gate obligatorio está bloqueado, no se activa el servicio. Toda excepción requiere decisión escrita de los socios y, cuando aplique, concepto profesional.", fill="FFF8E8", accent=GOLD)
    add_heading(doc, "8. Versiones y evidencias del expediente")
    add_matrix(doc, ("Documento / evidencia", "Versión", "Fecha", "Ubicación / enlace"), [
        ("Propuesta y cotización", "", "", ""),
        ("Contrato / términos", "", "", ""),
        ("Privacidad / transmisión", "", "", ""),
        ("Acta de Servicio", "", "", ""),
        ("Pago / soporte contable", "", "", ""),
        ("Capacitación y aceptación", "", "", ""),
    ], [2800, 1300, 1500, 3760])
    add_heading(doc, "9. Decisión final")
    add_key_values(doc, [
        ("Resultado", "☐ Activar  ☐ No activar  ☐ Activar con condición aprobada"),
        ("Fecha / hora efectiva", "____ / ____ / ______  ______  Zona __________________"),
        ("Condiciones / pendientes", "____________________________________________________________"),
        ("Próxima revisión", "____ / ____ / ______"),
    ])
    add_signature_table(doc, "Aprobación comercial Roustix", "Aprobación operativa Roustix")
    return save_document(doc, "19-CHECKLIST-CONTRATACION-ACTIVACION-ROUSTIX.docx", "Checklist Maestro de Contratación y Activación Roustix", "Control de expediente y activación")


def build_all() -> list[Path]:
    return [
        build_quote(),
        build_onboarding_guide(),
        build_training_act(),
        build_acceptance_act(),
        build_termination_act(),
        build_partner_decision_matrix(),
        build_review_pack(),
        build_contracting_checklist(),
    ]


if __name__ == "__main__":
    for generated in build_all():
        print(generated)
