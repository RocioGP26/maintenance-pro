"""Genera la propuesta comercial y el acta de servicio vigentes de Roustix."""

from __future__ import annotations

import os
from pathlib import Path

from docx import Document

from build_pilot_document_package import (
    BLACK,
    GOLD,
    NAVY,
    add_callout,
    add_heading,
    add_key_values,
    add_matrix,
    add_para,
    add_signature_table,
    start_document,
)


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = Path(
    os.environ.get(
        "ROUSTIX_COMMERCIAL_DOC_OUTPUT_DIR",
        ROOT / "docs" / "production-readiness" / "templates",
    )
)


def save_document(
    doc: Document,
    filename: str,
    title: str,
    subject: str,
    *,
    legal: bool = False,
) -> Path:
    props = doc.core_properties
    props.title = title
    props.subject = subject
    props.author = "Roustix"
    props.keywords = "Roustix, SaaS, propuesta comercial, activos, mantenimiento"
    props.comments = (
        "Borrador sujeto a revisión jurídica y tributaria."
        if legal
        else "Plantilla comercial corporativa para completar."
    )
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / filename
    doc.save(path)
    return path


def build_commercial_proposal() -> Path:
    doc = start_document(
        "Propuesta Comercial Roustix",
        "Gestión inteligente de activos · alcance, inversión y activación",
        "COM-PRO-001",
        "Plantilla comercial para completar",
        "Presentar al Cliente la solución propuesta, el plan, el alcance, la inversión y las condiciones necesarias para activar el servicio SaaS Roustix.",
        version="1.0",
    )

    add_heading(doc, "1. Cliente y oportunidad")
    add_key_values(doc, [
        ("Empresa / razón social", "____________________________________________________________"),
        ("NIT / identificación", "____________________________________________________________"),
        ("Contacto / cargo", "____________________________________________________________"),
        ("Sector / vertical", "☐ Manufactura  ☐ Educación  ☐ Salud  ☐ Hotelería  ☐ Transporte  ☐ Otro"),
        ("Necesidad prioritaria", "____________________________________________________________"),
        ("Fecha de presentación", "____ / ____ / ______"),
        ("Vigencia de la propuesta", "________ días calendario"),
    ])

    add_heading(doc, "2. Solución propuesta")
    add_para(
        doc,
        "Roustix centraliza la gestión de activos, mantenimiento, incidencias, órdenes de trabajo, jornadas, repuestos, costos, documentos y trazabilidad. La configuración se adapta al sector, los módulos contratados y los procesos autorizados por el Cliente.",
    )
    # Mantiene la introducción unida a la tabla y evita dejar una sola fila
    # del alcance funcional al final de una página.
    doc.paragraphs[-1].paragraph_format.keep_with_next = True
    add_matrix(doc, ("Componente", "Selección", "Resultado esperado"), [
        ("Gestión de activos y hoja de vida", "☐", "Inventario técnico, documentos, imágenes e historial consolidado."),
        ("Mantenimiento", "☐", "Preventivos, correctivos, OT, jornadas, técnicos, repuestos y costos."),
        ("Incidencias", "☐", "Reporte, asignación, diagnóstico, conversión a OT y cierre trazable."),
        ("Inventario y repuestos", "☐", "Existencias, entradas, proveedores y consumo en mantenimiento."),
        ("API y webhooks", "☐", "Integración según plan, alcance técnico y cotización aprobada."),
    ], [3000, 1100, 5260])
    solution_table = doc.tables[-1]
    for row in solution_table.rows[:2]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_with_next = True

    add_heading(doc, "3. Planes comerciales")
    add_matrix(doc, ("Plan", "Mensualidad", "Usuarios", "Sedes", "Módulos", "Almacenamiento", "Selección"), [
        ("Start", "$1.000.000 COP", "20", "1", "1 principal", "1 GB", "☐"),
        ("Business", "$1.500.000 COP", "50", "3", "Hasta 2", "5 GB", "☐"),
        ("Enterprise", "Según cotización", "Personalizado", "Personalizadas", "Todos / alcance", "20 GB ampliables", "☐"),
    ], [1150, 1750, 1050, 900, 1650, 1700, 1160], size=8.6)
    add_callout(
        doc,
        "Referencia comercial",
        "Start y Business usan los precios de lista vigentes. Enterprise se cotiza según módulos, capacidad, integraciones, soporte y nivel de servicio. Los impuestos aplicables se confirman antes de la aceptación definitiva.",
        fill="FFF8E8",
        accent=GOLD,
    )

    add_heading(doc, "4. Alcance incluido")
    add_matrix(doc, ("Componente", "Condición del servicio"), [
        ("Activación", "Creación del tenant y configuración remota del alcance contratado."),
        ("Usuarios, sedes y módulos", "Según el plan seleccionado y las condiciones especiales aprobadas."),
        ("Activos", "Ilimitados dentro del uso razonable y la capacidad técnica contratada."),
        ("Seguridad", "HTTPS, autenticación, roles, aislamiento entre empresas y auditoría."),
        ("Archivos", "Almacenamiento persistente sujeto a la cuota efectiva del plan."),
        ("Respaldos", "Controles de respaldo y recuperación conforme a las condiciones vigentes."),
        ("Soporte", "Email para Start; chat para Business; canal dedicado según Orden para Enterprise."),
        ("Capacitación inicial", "Sesión remota para administradores y responsables según alcance."),
        ("Actualizaciones", "Correcciones y mejoras generales durante la vigencia del servicio."),
    ], [2800, 6560])

    add_heading(doc, "5. Complementos y servicios opcionales")
    add_matrix(doc, ("Complemento", "Referencia", "Valor"), [
        ("Almacenamiento +2 GB", "ADD-STG-2G", "$100.000 COP / mes"),
        ("Almacenamiento +5 GB", "ADD-STG-5G", "$220.000 COP / mes"),
        ("Almacenamiento +10 GB", "ADD-STG-10G", "$400.000 COP / mes"),
        ("Usuarios adicionales", "ADD-USR-COT", "Según cotización"),
        ("Capacitación adicional", "ADD-TRN-COT", "Según cotización"),
        ("Implementación en sitio", "ADD-ONB-SITE", "Según cotización"),
        ("Integraciones especiales", "ADD-INT-COT", "Según cotización"),
    ], [3800, 2200, 3360])

    add_heading(doc, "6. Ruta de activación")
    add_matrix(doc, ("Etapa", "Responsable", "Entregable / criterio"), [
        ("Aceptación comercial", "Ambas partes", "Propuesta, contrato y acta de servicio aceptados."),
        ("Preparación", "Cliente", "Responsables, datos iniciales y alcance confirmados."),
        ("Configuración", "Roustix", "Tenant, plan, sedes, módulos y permisos configurados."),
        ("Capacitación", "Ambas partes", "Administradores y perfiles clave orientados."),
        ("Activación", "Roustix", "Gate de activación aprobado y servicio habilitado."),
        ("Seguimiento", "Ambas partes", "Revisión de adopción en la fecha acordada."),
    ], [2400, 2200, 4760])

    add_heading(doc, "7. Resumen económico")
    add_key_values(doc, [
        ("Plan seleccionado", "☐ Start  ☐ Business  ☐ Enterprise  ☐ Personalizado"),
        ("Mensualidad base", "$ __________________________ COP"),
        ("Complementos mensuales", "$ __________________________ COP"),
        ("Servicios de pago único", "$ __________________________ COP"),
        ("Descuento / condición especial", "____________________________________________________________"),
        ("Impuestos aplicables", "____________________________________________________________"),
        ("Total periódico", "$ __________________________ COP"),
        ("Periodicidad", "☐ Mensual anticipada  ☐ Otra: ______________________________"),
        ("Fecha estimada de activación", "____ / ____ / ______"),
    ])

    # El separador largo evita una colisión tipográfica observada al exportar
    # el par "8. C" con Word en algunos entornos.
    add_heading(doc, "8 — Condiciones de la propuesta")
    add_para(
        doc,
        "La contratación queda integrada por la propuesta aceptada, el Contrato SaaS, el Acta de Servicio y Selección de Plan, los documentos de privacidad y transmisión aplicables, la Política de Soporte y el SLA únicamente cuando se anexe expresamente.",
    )
    add_callout(
        doc,
        "Validación previa a la firma",
        "La identidad de la parte prestadora, la facturación, los impuestos, la vigencia, la renovación, la terminación y cualquier condición especial deben quedar completos y revisados antes de emitir cobros o activar el servicio.",
        fill="FFF8E8",
        accent=GOLD,
    )

    add_heading(doc, "9. Aceptación comercial")
    add_para(
        doc,
        "La firma confirma la aceptación de esta propuesta y autoriza la preparación del Acta de Servicio. La activación solo procede cuando se complete el gate contractual, operativo y de pago aplicable.",
    )
    add_signature_table(doc, "Empresa Cliente", "Prestador Roustix")

    return save_document(
        doc,
        "01-PROPUESTA-COMERCIAL-ROUSTIX.docx",
        "Propuesta Comercial Roustix",
        "Alcance, inversión y activación del servicio SaaS",
    )


def build_service_and_plan_act() -> Path:
    doc = start_document(
        "Acta de Servicio y Selección de Plan",
        "Orden de activación, alcance contratado y responsables",
        "SAA-ACT-001",
        "Plantilla comercial para completar y firmar",
        "Dejar constancia del servicio Roustix contratado por el Cliente, el plan, las capacidades, complementos, valores, vigencia, responsables y condiciones de activación.",
        version="2.0",
    )
    add_callout(
        doc,
        "Documento contractual complementario",
        "Esta acta concreta la Orden de servicio y debe leerse junto con la propuesta aceptada, el Contrato SaaS, los documentos de privacidad y transmisión, la Política de Soporte y el SLA cuando se anexe expresamente.",
        fill="FFF8E8",
        accent=GOLD,
    )

    add_heading(doc, "1. Identificación del servicio")
    add_key_values(doc, [
        ("Empresa / razón social", "____________________________________________________________"),
        ("NIT / identificación", "____________________________________________________________"),
        ("Representante autorizado", "____________________________________________________________"),
        ("Contacto administrativo", "____________________________________________________________"),
        ("Contacto de facturación", "____________________________________________________________"),
        ("Tenant / slug", "____________________________________________________________"),
        ("Sector / vertical", "☐ Manufactura  ☐ Educación  ☐ Salud  ☐ Hotelería  ☐ Transporte  ☐ Otro"),
        ("Tipo de Orden", "☐ Nuevo servicio  ☐ Renovación  ☐ Ampliación  ☐ Modificación"),
    ])

    add_heading(doc, "2. Planes comerciales Roustix")
    add_matrix(doc, ("Plan", "Mensualidad", "Usuarios", "Sedes", "Módulos principales", "Almacenamiento", "Selección"), [
        ("Start", "$1.000.000 COP", "20", "1", "1 a elegir", "1 GB", "☐"),
        ("Business", "$1.500.000 COP", "50", "3", "Hasta 2", "5 GB", "☐"),
        ("Enterprise", "Según cotización", "Personalizado", "Personalizadas", "Todos / alcance", "20 GB ampliables", "☐"),
    ], [1200, 1800, 1050, 950, 1850, 1500, 1010], size=8.7)
    add_para(doc, "Enterprise se cotiza de acuerdo con el alcance definitivo. Los valores, impuestos, descuentos y condiciones especiales deben quedar diligenciados en la sección económica de esta acta.")

    add_heading(doc, "3. Plan y alcance seleccionado")
    add_key_values(doc, [
        ("Plan seleccionado", "☐ Start  ☐ Business  ☐ Enterprise  ☐ Personalizado"),
        ("Número de usuarios habilitados", "________________"),
        ("Número de sedes habilitadas", "________________"),
        ("Cuota base de almacenamiento", "________________ GB"),
        ("Módulo principal 1", "____________________________________________________________"),
        ("Módulo principal 2", "____________________________________________________________"),
        ("Otros módulos Enterprise", "____________________________________________________________"),
        ("Fecha prevista de activación", "____ / ____ / ______"),
    ])

    add_heading(doc, "4. Capacidades incluidas")
    add_matrix(doc, ("Capacidad", "Condición del servicio"), [
        ("Activos", "Ilimitados dentro del uso razonable, módulos habilitados y capacidad técnica acordada."),
        ("Seguridad", "HTTPS, autenticación, permisos por rol, aislamiento entre empresas y auditoría."),
        ("Archivos", "Almacenamiento persistente sujeto a la cuota efectiva del plan y complementos."),
        ("Respaldo", "Respaldos y recuperación conforme a los controles y condiciones vigentes."),
        ("Actualizaciones", "Correcciones y mejoras generales de la plataforma durante la vigencia."),
        ("Soporte", "Email para Start; chat para Business; dedicado y SLA acordado para Enterprise."),
        ("Onboarding inicial", "Activación del tenant y configuración remota guiada del alcance contratado."),
        ("Capacitación inicial", "Sesión remota para administradores y responsables, según la propuesta aceptada."),
        ("Seguimiento de adopción", "Revisión inicial de uso y hallazgos dentro del periodo acordado."),
        ("Acceso multidispositivo", "Mediante navegador compatible y cuentas personales autorizadas."),
    ], [2600, 6760])

    add_heading(doc, "5. Complementos recurrentes")
    add_matrix(doc, ("Complemento", "SKU", "Valor mensual", "Cantidad", "Seleccionado"), [
        ("Almacenamiento +2 GB", "ADD-STG-2G", "$100.000 COP", "______", "☐"),
        ("Almacenamiento +5 GB", "ADD-STG-5G", "$220.000 COP", "______", "☐"),
        ("Almacenamiento +10 GB", "ADD-STG-10G", "$400.000 COP", "______", "☐"),
        ("Usuarios adicionales", "ADD-USR-COT", "Según cotización", "______", "☐"),
    ], [2600, 1600, 2000, 1300, 1860], size=8.9)
    add_para(doc, "La activación de complementos sujetos a pago se realiza después de confirmar la condición comercial acordada. El retiro de capacidad no elimina automáticamente archivos existentes.")

    add_heading(doc, "6. Servicios profesionales opcionales")
    add_callout(
        doc,
        "Valores de referencia",
        "Los valores se confirman en la propuesta o cotización aceptada. El alcance, impuestos, desplazamientos, dependencias y entregables deben quedar por escrito antes de iniciar.",
        fill="FFF8E8",
        accent=GOLD,
    )
    add_matrix(doc, ("Servicio", "SKU", "Unidad / alcance de referencia", "Valor", "Selección"), [
        ("Capacitación adicional", "ADD-TRN-COT", "Sesión o jornada acordada", "Según cotización", "☐"),
        ("Implementación en sitio", "ADD-ONB-SITE", "Jornada y ubicación acordadas", "Según cotización", "☐"),
        ("Migración inicial de activos", "ADD-MIG-COT", "Plantilla y volumen validados", "Según cotización", "☐"),
        ("Depuración y organización de datos", "ADD-DATA-COT", "Alcance por diagnóstico", "Según cotización", "☐"),
        ("Formato o informe personalizado", "ADD-RPT-COT", "Por formato aprobado", "Según cotización", "☐"),
        ("Integración API / webhooks", "ADD-INT-COT", "Por integración y alcance técnico", "Según cotización", "☐"),
        ("Desarrollo especial", "ADD-DEV-COT", "Bolsa o proyecto estimado", "Según cotización", "☐"),
    ], [2200, 1450, 2800, 1800, 1110], size=8.35)
    add_key_values(doc, [
        ("Servicio profesional seleccionado", "____________________________________________________________"),
        ("Alcance / entregable", "____________________________________________________________"),
        ("Fecha o plazo estimado", "____________________________________________________________"),
        ("Valor definitivo aprobado", "$ __________________________ COP"),
    ])

    add_heading(doc, "7. Condiciones económicas")
    add_key_values(doc, [
        ("Mensualidad base", "$ __________________________ COP"),
        ("Complementos mensuales", "$ __________________________ COP"),
        ("Servicios de pago único", "$ __________________________ COP"),
        ("Descuento / condición especial", "____________________________________________________________"),
        ("Subtotal", "$ __________________________ COP"),
        ("Impuestos aplicables", "____________________________________________________________"),
        ("Total periódico", "$ __________________________ COP"),
        ("Periodicidad", "☐ Mensual anticipada  ☐ Otra: ______________________________"),
        ("Medio de pago", "____________________________________________________________"),
        ("Fecha de vencimiento", "____________________________________________________________"),
    ])
    add_callout(doc, "Validación tributaria", "La parte prestadora debe validar facturación e impuestos aplicables según su situación jurídica y tributaria antes de emitir cobros o firmar la versión definitiva.", fill="FFF8E8", accent=GOLD)

    add_heading(doc, "8. Vigencia, renovación y terminación")
    add_key_values(doc, [
        ("Inicio del servicio", "____ / ____ / ______"),
        ("Vigencia inicial", "________ meses"),
        ("Renovación", "☐ Automática  ☐ Por acuerdo expreso  ☐ Sin renovación automática"),
        ("Preaviso de no renovación", "________ días calendario"),
        ("Plazo de exportación al terminar", "________ días"),
        ("Conservación residual de respaldos", "________ días / según política aprobada"),
    ])

    add_heading(doc, "9. Responsables y soporte")
    add_matrix(doc, ("Rol", "Nombre", "Correo / teléfono", "Responsabilidad"), [
        ("Administrador del Cliente", "", "", "Usuarios, permisos y configuración"),
        ("Responsable operativo", "", "", "Procesos, activos y mantenimiento"),
        ("Facturación del Cliente", "", "", "Pagos y documentos comerciales"),
        ("Onboarding Roustix", "", "", "Activación y acompañamiento"),
        ("Soporte Roustix", "", "", "Incidentes y escalamiento"),
    ], [1900, 1900, 2500, 3060])
    add_key_values(doc, [
        ("Canal de soporte", "contacto@roustix.com"),
        ("Horario acordado", "Lunes a viernes: ______ a ______ · America/Bogota"),
        ("Canal alterno pactado", "____________________________________________________________"),
    ])

    add_heading(doc, "10. Gate de activación")
    add_matrix(doc, ("Verificación", "Estado", "Evidencia / observación"), [
        ("Empresa, plan y tenant correctamente identificados", "☐", ""),
        ("Usuarios, sedes y módulos dentro del alcance", "☐", ""),
        ("Administrador y responsables designados", "☐", ""),
        ("Propuesta, contrato y acta aceptados", "☐", ""),
        ("Privacidad y transmisión entregadas / aceptadas", "☐", ""),
        ("Política de soporte entregada", "☐", ""),
        ("Condición de pago validada", "☐", ""),
        ("Fecha y responsable de activación confirmados", "☐", ""),
    ], [5100, 900, 3360])

    add_heading(doc, "11. Documentos que integran el servicio")
    add_matrix(doc, ("Documento", "Código / versión", "Fecha / aceptación"), [
        ("Propuesta comercial", "COM-PRO-001", ""),
        ("Contrato SaaS", "RTX-LEGAL-002", ""),
        ("Acta de Servicio y Selección de Plan", "SAA-ACT-001", ""),
        ("Política de privacidad y anexos", "RTX-PRIV-001", ""),
        ("Acuerdo de transmisión de datos", "PIL-LEG-003 / versión vigente", ""),
        ("Política de soporte", "RTX-SUP-001", ""),
        ("SLA, si se anexa", "RTX-SLA-001", ""),
        ("Otros anexos", "", ""),
    ], [4200, 2400, 2760])

    add_heading(doc, "12. Aceptación y autorización de activación")
    add_para(doc, "Las partes confirman que el alcance, plan, capacidades, complementos, servicios profesionales, valores y responsables corresponden a lo acordado. La firma autoriza la activación del servicio en la fecha indicada, sin reemplazar las obligaciones de los demás documentos integrantes.")
    add_signature_table(doc, "Empresa Cliente", "Prestador Roustix")

    return save_document(
        doc,
        "09-ACTA-SERVICIO-SELECCION-PLAN-ROUSTIX.docx",
        "Acta de Servicio y Selección de Plan Roustix",
        "Orden de activación y alcance del servicio SaaS",
        legal=True,
    )


def build_all() -> list[Path]:
    return [build_commercial_proposal(), build_service_and_plan_act()]


if __name__ == "__main__":
    for generated in build_all():
        print(generated)
