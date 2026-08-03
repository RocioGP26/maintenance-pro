"""Genera el paquete documental mínimo para pilotos de Roustix."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from build_pilot_linkage_act import (
    BLACK,
    BLUE,
    CORPORATE_LOGO,
    GOLD,
    GRAY,
    LIGHT_BLUE,
    NAVY,
    PALE_BLUE,
    RED,
    add_callout,
    add_cell_text,
    add_heading,
    add_word_field,
    configure_styles,
    enable_field_updates,
    set_cell_shading,
    set_run_font,
    set_table_geometry,
    style_table,
)


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs" / "production-readiness" / "templates"
WHITE = "FFFFFF"


def configure_page(doc: Document, footer_label: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("ROUSTIX  ·  PAQUETE DOCUMENTAL")
    set_run_font(r, size=8.5, color=GRAY, bold=True)

    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    add_cell_text(table.cell(0, 0), footer_label, color=GRAY, size=8.2)
    p = table.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Página ")
    set_run_font(r, size=8.2, color=GRAY)
    add_word_field(p, "PAGE")
    r = p.add_run(" de ")
    set_run_font(r, size=8.2, color=GRAY)
    add_word_field(p, "NUMPAGES")
    set_table_geometry(table, [6500, 2860], indent=0)


def start_document(
    title: str,
    subtitle: str,
    code: str,
    status: str,
    purpose: str,
    version: str = "1.0",
) -> Document:
    doc = Document()
    configure_styles(doc)
    configure_page(doc, f"{code} · Versión {version}")
    enable_field_updates(doc)

    if not CORPORATE_LOGO.is_file():
        raise FileNotFoundError(f"No se encontró el logotipo corporativo: {CORPORATE_LOGO}")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    shape = p.add_run().add_picture(str(CORPORATE_LOGO), width=Inches(2.6))
    shape._inline.docPr.set("title", "Roustix")
    shape._inline.docPr.set("descr", "Logotipo corporativo de Roustix")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(title)
    set_run_font(r, size=25, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(13)
    r = p.add_run(subtitle)
    set_run_font(r, size=12.3, color=GRAY)

    table = doc.add_table(rows=2, cols=4)
    values = [
        ("Código", code, "Versión", version),
        ("Estado", status, "Fecha", "____ / ____ / ______"),
    ]
    for row, data in zip(table.rows, values):
        for index, value in enumerate(data):
            add_cell_text(row.cells[index], value, bold=index % 2 == 0, color=NAVY if index % 2 == 0 else BLACK)
            if index % 2 == 0:
                set_cell_shading(row.cells[index], LIGHT_BLUE)
    set_table_geometry(table, [1300, 3000, 1300, 3760])
    style_table(table, header=False, label_columns={0, 2})
    doc.add_paragraph()
    add_callout(doc, "Propósito", purpose, fill=PALE_BLUE, accent=BLUE)
    return doc


def add_para(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.keep_together = True
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run_font(r, size=11, color=NAVY, bold=True)
        r = p.add_run(text[len(bold_lead):])
        set_run_font(r, size=11, color=BLACK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=BLACK)


def add_key_values(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    for row, (label, value) in zip(table.rows, rows):
        add_cell_text(row.cells[0], label, bold=True, color=NAVY)
        add_cell_text(row.cells[1], value)
        set_cell_shading(row.cells[0], LIGHT_BLUE)
    set_table_geometry(table, [2800, 6560])
    style_table(table, header=False, label_columns={0})


def add_matrix(doc: Document, headers: tuple[str, ...], rows: list[tuple[str, ...]], widths: list[int], size=9.4) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for cell, value in zip(table.rows[0].cells, headers):
        add_cell_text(cell, value, bold=True, color=WHITE, size=size)
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            add_cell_text(cell, value, size=size)
    set_table_geometry(table, widths)
    style_table(table, header=True)


def add_party_block(doc: Document, include_two_operators=True) -> None:
    rows = [
        ("Cliente / razón social", "____________________________________________________________"),
        ("NIT / identificación", "____________________________________________________________"),
        ("Representante autorizado", "____________________________________________________________"),
        ("Domicilio y contacto", "____________________________________________________________"),
    ]
    if include_two_operators:
        rows.extend([
            ("Operador 1 de Roustix", "Nombre, identificación, domicilio y contacto: ______________________________"),
            ("Operador 2 de Roustix", "Nombre, identificación, domicilio y contacto: ______________________________"),
        ])
    add_key_values(doc, rows)


def add_signature_table(doc: Document, left="Cliente", right="Roustix") -> None:
    table = doc.add_table(rows=5, cols=2)
    values = [
        (left, right),
        ("Nombre: __________________________", "Nombre: __________________________"),
        ("Cargo / calidad: __________________", "Cargo / calidad: __________________"),
        ("Firma: ___________________________", "Firma: ___________________________"),
        ("Fecha: ____ / ____ / ______", "Fecha: ____ / ____ / ______"),
    ]
    for row, data in zip(table.rows, values):
        for index, value in enumerate(data):
            add_cell_text(row.cells[index], value, bold=row._index == 0, color=NAVY if row._index == 0 else BLACK)
            if row._index == 0:
                set_cell_shading(row.cells[index], LIGHT_BLUE)
    set_table_geometry(table, [4680, 4680])
    style_table(table, header=False)


def save_document(doc: Document, filename: str, title: str, subject: str, legal=False) -> Path:
    props = doc.core_properties
    props.title = title
    props.subject = subject
    props.author = "Roustix"
    props.keywords = "Roustix, piloto, SaaS, documento empresarial"
    props.comments = "Borrador sujeto a revisión jurídica." if legal else "Plantilla corporativa para completar."
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / filename
    doc.save(path)
    return path


def build_commercial_proposal() -> Path:
    doc = start_document(
        "Propuesta Comercial del Piloto",
        "Alcance, inversión y condiciones para una incorporación controlada",
        "PIL-COM-001",
        "Plantilla para completar",
        "Presentar en un solo documento qué probará la empresa, durante cuánto tiempo, con qué capacidad, acompañamiento y condiciones económicas.",
    )
    add_heading(doc, "1. Empresa y oportunidad")
    add_key_values(doc, [
        ("Empresa", "____________________________________________________________"),
        ("Contacto / cargo", "____________________________________________________________"),
        ("Sector y operación", "____________________________________________________________"),
        ("Necesidad prioritaria", "____________________________________________________________"),
        ("Fecha de presentación", "____ / ____ / ______"),
        ("Vigencia de la propuesta", "________ días calendario"),
    ])
    add_heading(doc, "2. Objetivo del piloto")
    add_para(doc, "Validar, con usuarios y datos controlados de la empresa, que Roustix permite gestionar activos, incidencias, órdenes de trabajo, jornadas, repuestos, documentos y seguimiento operativo según los módulos habilitados.")
    add_heading(doc, "3. Alcance ofrecido")
    add_matrix(doc, ("Componente", "Incluido", "Condición / límite"), [
        ("Onboarding asistido", "Sí", "Configuración inicial de empresa, sede, áreas y responsables."),
        ("Usuarios", "________", "Máximo autorizado durante el piloto."),
        ("Módulos", "________________", "Solo los expresamente habilitados."),
        ("Activos", "________________", "Carga inicial acordada con la empresa."),
        ("Almacenamiento", "________ GB", "Plan base más add-on aprobado, si aplica."),
        ("Soporte", "________________", "Canal y horario definidos en la Guía Operativa."),
        ("UAT", "Sí", "Recorrido abreviado y evidencia de aceptación."),
    ], [2500, 1500, 5360])
    add_heading(doc, "4. Plan e inversión")
    add_matrix(doc, ("Plan", "Referencia mensual", "Almacenamiento", "Selección"), [
        ("Start", "$1.000.000 COP", "1 GB", "☐"),
        ("Business", "$1.500.000 COP", "5 GB", "☐"),
        ("Enterprise", "Según alcance", "20 GB ampliables", "☐"),
        ("Piloto especial", "$ __________________", "________ GB", "☐"),
    ], [2100, 2600, 2500, 2160])
    add_para(doc, "Los precios, impuestos, gratuidad, descuentos y forma de pago solo quedan acordados cuando se completan y firman la propuesta y el Acta de Vinculación.")
    add_heading(doc, "5. Cronograma")
    add_matrix(doc, ("Etapa", "Fecha", "Responsable", "Entregable"), [
        ("Preparación", "____________", "Ambas partes", "Datos y responsables definidos"),
        ("Configuración", "____________", "Roustix", "Tenant habilitado"),
        ("Ejecución UAT", "____________", "Empresa", "Evidencia funcional"),
        ("Seguimiento", "____________", "Ambas partes", "Hallazgos y acciones"),
        ("Cierre", "____________", "Ambas partes", "Informe y decisión"),
    ], [2200, 1700, 2200, 3260])
    add_heading(doc, "6. Condiciones de aceptación")
    add_para(doc, "La propuesta se integra con los Términos del Piloto, el Acta de Vinculación, la Guía Operativa y los instrumentos de privacidad y transmisión aplicables.")
    add_signature_table(doc, "Empresa piloto", "Operadores de Roustix")
    return save_document(doc, "01-PROPUESTA-COMERCIAL-PILOTO-ROUSTIX.docx", "Propuesta Comercial del Piloto Roustix", "Alcance e inversión del piloto")


def build_pilot_terms() -> Path:
    doc = start_document(
        "Términos del Piloto SaaS",
        "Condiciones provisionales de acceso y uso empresarial",
        "PIL-LEG-001",
        "Borrador sujeto a revisión jurídica",
        "Regular el acceso temporal y asistido a Roustix. Debe completarse con la identificación de los operadores y revisarse jurídicamente antes de firmarse.",
    )
    add_callout(doc, "Advertencia", "Roustix es actualmente un nombre comercial operado por dos personas naturales. Si se constituye una sociedad antes de la firma, debe sustituirse íntegramente la identificación de la parte prestadora.", fill="FFF8E8", accent=GOLD)
    sections = [
        ("1. Partes y documentos integrantes", "El Cliente Piloto y los operadores de Roustix se identifican en el Acta de Vinculación. Integran el acuerdo la propuesta aceptada, esta versión de los términos, la guía operativa y los anexos de datos aplicables."),
        ("2. Objeto", "El piloto permite validar temporalmente funciones de gestión de activos, mantenimiento, inventario y módulos habilitados. No concede propiedad sobre el software ni autoriza su reventa, copia, ingeniería inversa o uso por terceros no vinculados."),
        ("3. Alcance económico", "Plan, precio o gratuidad, impuestos, forma de pago, usuarios, módulos, duración y almacenamiento deben constar en la propuesta y el acta. Ningún cobro o beneficio adicional se presume."),
        ("4. Cuentas y uso autorizado", "El Cliente designará administradores, mantendrá datos exactos, asignará permisos mínimos, protegerá credenciales, dará de baja accesos oportunamente y no cargará contenido ilícito, malicioso o sin autorización."),
        ("5. Datos del cliente", "El Cliente conserva control sobre los datos que incorpora. Para esos datos determina normalmente las finalidades y Roustix actúa como encargado conforme al acuerdo de transmisión y a instrucciones documentadas."),
        ("6. Propiedad intelectual", "Los operadores conservan los derechos sobre software, diseño, documentación, marca y mejoras. El Cliente conserva sus marcas, documentos y datos. La retroalimentación podrá usarse para mejorar el producto sin identificar al Cliente ni revelar información confidencial."),
        ("7. Disponibilidad y cambios", "La versión piloto puede requerir mantenimiento y correcciones. Interrupciones relevantes se comunicarán cuando sea razonablemente posible. RTO, RPO, disponibilidad, créditos o compensaciones solo constituyen SLA si se pactan expresamente."),
        ("8. Soporte", "El canal, horario, severidades y procedimiento de reporte serán los definidos en la Guía Operativa y el Acta de Vinculación. El Cliente no incluirá contraseñas, tokens o datos innecesarios en solicitudes."),
        ("9. Seguridad y continuidad", "Roustix mantiene controles de autenticación, permisos, aislamiento por empresa, cifrado en tránsito, auditoría, almacenamiento persistente, respaldos, monitoreo y recuperación. El Cliente responde por dispositivos, usuarios, permisos y exactitud de la información."),
        ("10. Confidencialidad", "Cada parte protegerá la información no pública recibida, la usará solo para el piloto y limitará su acceso a quienes necesiten conocerla. La obligación continuará después del cierre mientras la información conserve carácter confidencial."),
        ("11. Suspensión y terminación", "El acceso podrá suspenderse para contener riesgos, uso ilícito, incumplimiento material o impago expresamente pactado, procurando aviso y oportunidad de corrección cuando sea posible. La duración y causales particulares se completan en el acta."),
        ("12. Exportación y eliminación", "Al terminar se aplicará el plazo de exportación y la matriz de conservación expresamente aprobados. Las copias activas se eliminarán o anonimizarán cuando corresponda y los respaldos se depurarán por rotación, salvo obligación legal."),
        ("13. Responsabilidad", "Cada parte responde por sus obligaciones y por los daños legalmente imputables. Cualquier límite monetario, exclusión, indemnidad o distribución especial de riesgos debe negociarse expresamente y revisarse jurídicamente."),
        ("14. Ley, controversias y cambios", "Ley aplicable, jurisdicción y mecanismo de solución se completarán con asesoría jurídica. Los cambios sustanciales se informarán y la versión aplicable será la aceptada por las partes."),
    ]
    for heading, text in sections:
        add_heading(doc, heading)
        add_para(doc, text)
    add_heading(doc, "15. Datos pendientes antes de firma")
    add_key_values(doc, [
        ("Identificación completa de ambos operadores", "PENDIENTE"),
        ("Canal y horario de soporte", "PENDIENTE DE APROBACIÓN"),
        ("Plazo de exportación", "________ días"),
        ("Rotación residual de respaldos", "________ días"),
        ("Ley / jurisdicción / conciliación", "PENDIENTE DE REVISIÓN JURÍDICA"),
        ("Tratamiento tributario", "PENDIENTE DE REVISIÓN CONTABLE"),
    ])
    add_heading(doc, "16. Aceptación")
    add_signature_table(doc, "Cliente Piloto", "Operadores de Roustix")
    return save_document(doc, "02-TERMINOS-PILOTO-SAAS-ROUSTIX.docx", "Términos del Piloto SaaS Roustix", "Condiciones provisionales del piloto", legal=True)


def build_privacy_package() -> Path:
    doc = start_document(
        "Paquete de Privacidad y Tratamiento",
        "Política, aviso de privacidad y formato de autorización",
        "PIL-LEG-002",
        "Borrador sujeto a revisión jurídica",
        "Reunir en un único archivo los textos que deben informarse y, cuando corresponda, aceptarse antes o al momento de recolectar datos personales.",
    )
    add_callout(doc, "Pendiente crítico", "Completar la identificación de ambos responsables, el canal definitivo de titulares, la lista de subencargados, los flujos internacionales y la matriz de conservación antes de publicar.", fill="FCEDED", accent=RED)
    add_heading(doc, "Parte A. Política de Tratamiento y Protección")
    policy = [
        ("1. Responsable", "Mientras se constituye la sociedad, los responsables serán las personas naturales operadoras identificadas en la versión firmada. Nombre comercial: Roustix. Sitio: https://roustix.com. Canal provisional: soporte.roustix@hotmail.com."),
        ("2. Alcance", "Aplica a datos recolectados mediante registro, plataforma SaaS, soporte, comunicaciones, facturación y piloto. En los datos cargados por una empresa cliente, esta normalmente actúa como responsable y Roustix como encargado."),
        ("3. Categorías", "Identificación y contacto; cargo, área y empresa; autenticación, sesiones, IP y auditoría; información contractual y de soporte; registros operativos; fotografías, informes, evidencias y datos técnicos."),
        ("4. Finalidades", "Crear y proteger cuentas; prestar módulos; administrar permisos; enviar comunicaciones necesarias; atender soporte y derechos; respaldar y recuperar; auditar; prevenir abuso; mejorar seguridad y rendimiento; cumplir obligaciones legales."),
        ("5. Autorización", "Cuando sea necesaria será previa, expresa e informada y se conservará evidencia de versión, fecha, identidad y medio. Una casilla electrónica no estará premarcada."),
        ("6. Derechos", "Conocer, actualizar, rectificar y acceder; solicitar prueba de autorización; conocer el uso; presentar consultas o reclamos; revocar o solicitar supresión cuando proceda; acudir a la SIC después del trámite aplicable."),
        ("7. Consultas y reclamos", "Se verificará la identidad antes de revelar o modificar datos. Consultas: hasta 10 días hábiles, prorrogables por 5; reclamos: hasta 15 días hábiles, prorrogables por 8, conforme al régimen aplicable."),
        ("8. Seguridad", "Se aplican controles proporcionales de acceso, aislamiento por tenant, cifrado en tránsito, credenciales protegidas, almacenamiento de objetos, auditoría, respaldos, monitoreo y gestión de incidentes."),
        ("9. Proveedores", "Pueden intervenir proveedores de alojamiento, base de datos, almacenamiento, correo y observabilidad. La versión final identificará entidad, servicio, región, datos tratados y mecanismo de información de cambios."),
        ("10. Conservación", "Los datos se conservarán durante la relación y por plazos necesarios para finalidades, obligaciones legales, seguridad, respaldo o defensa de derechos. Después se eliminarán o anonimizarán de manera segura."),
        ("11. Cookies", "Solo se contemplan cookies necesarias para sesión, autenticación, seguridad y preferencias. Analítica o publicidad no esencial requerirá evaluación y, cuando corresponda, consentimiento."),
        ("12. Vigencia", "La versión final indicará fecha de entrada en vigor y se comunicará cualquier cambio sustancial. Si se constituye la sociedad operadora, se actualizará la identidad del responsable."),
    ]
    for heading, text in policy:
        add_heading(doc, heading, level=2)
        add_para(doc, text)

    add_heading(doc, "Parte B. Aviso de Privacidad")
    add_callout(doc, "Texto corto para formularios", "Roustix tratará los datos suministrados para registrar y proteger la cuenta, prestar la plataforma, gestionar usuarios y permisos, enviar comunicaciones necesarias, brindar soporte, mantener seguridad y auditoría y administrar la relación contractual. El titular puede ejercer sus derechos mediante el canal informado. La política completa estará disponible en: ________________________________.", fill=PALE_BLUE, accent=BLUE)

    add_heading(doc, "Parte C. Autorización del titular")
    add_para(doc, "Declaro que recibí información clara sobre el responsable, las finalidades, el carácter facultativo de respuestas sobre datos sensibles, mis derechos y el canal para ejercerlos. Cuando la autorización sea necesaria, autorizo el tratamiento descrito en la versión identificada a continuación.")
    add_key_values(doc, [
        ("Titular", "____________________________________________________________"),
        ("Identificación", "____________________________________________________________"),
        ("Correo / contacto", "____________________________________________________________"),
        ("Finalidades aceptadas", "☐ Cuenta y servicio  ☐ Soporte  ☐ Seguridad  ☐ Comunicaciones comerciales separadas"),
        ("Versión de política y aviso", "____________________________________________________________"),
        ("Medio de autorización", "☐ Firma  ☐ Casilla electrónica no premarcada  ☐ Otro: ______________"),
        ("Fecha, hora y evidencia", "____________________________________________________________"),
    ])
    add_para(doc, "La autorización para comunicaciones comerciales debe ser separable y revocable; la negativa no debe impedir funciones esenciales que no dependan de esa finalidad.")
    add_heading(doc, "Referencias normativas para revisión")
    add_para(doc, "Constitución Política de Colombia, artículo 15; Ley Estatutaria 1581 de 2012; Decreto 1074 de 2015 y demás normas aplicables. La versión definitiva debe ser revisada por asesor jurídico colombiano.")
    add_signature_table(doc, "Titular / representante", "Responsable que recibe")
    return save_document(doc, "03-PAQUETE-PRIVACIDAD-TRATAMIENTO-ROUSTIX.docx", "Paquete de Privacidad y Tratamiento Roustix", "Política, aviso y autorización", legal=True)


def build_data_processing_agreement() -> Path:
    doc = start_document(
        "Acuerdo de Transmisión de Datos",
        "Anexo de encargo entre la empresa cliente y los operadores de Roustix",
        "PIL-LEG-003",
        "Borrador sujeto a revisión jurídica",
        "Documentar las instrucciones bajo las cuales Roustix trata datos personales por cuenta de la empresa durante el piloto o servicio.",
    )
    add_heading(doc, "1. Partes")
    add_party_block(doc)
    sections = [
        ("2. Objeto y duración", "El Cliente, como Responsable respecto de los datos que incorpora, encarga el tratamiento necesario para alojar, organizar, consultar, respaldar, proteger, recuperar, exportar y suprimir información en los módulos habilitados. El encargo dura durante el servicio y los períodos residuales definidos por escrito."),
        ("3. Instrucciones", "Roustix tratará los datos conforme al acuerdo, la configuración del Cliente, las acciones de usuarios autorizados y otras instrucciones escritas y lícitas. No determinará finalidades incompatibles ni venderá datos."),
        ("4. Titulares y datos", "Pueden comprender administradores, empleados, técnicos, solicitantes, responsables de activos, contactos de proveedores y terceros registrados; y datos de identificación, contacto, cargo, registros operativos, auditoría, archivos y evidencias."),
        ("5. Datos especiales", "No se autoriza cargar datos sensibles o de menores sin anexo previo, base jurídica demostrada, evaluación de riesgos y controles adicionales aprobados."),
    ]
    for heading, text in sections:
        add_heading(doc, heading)
        add_para(doc, text)
    add_heading(doc, "6. Obligaciones de Roustix")
    add_matrix(doc, ("Obligación", "Aplicación"), [
        ("Instrucciones y confidencialidad", "Tratar solo para finalidades documentadas y limitar el acceso a personal autorizado."),
        ("Seguridad", "Mantener controles proporcionales, aislamiento por tenant, trazabilidad y continuidad."),
        ("Derechos", "Cooperar razonablemente con consultas, reclamos, acceso, corrección, exportación y supresión."),
        ("Incidentes", "Informar sin demora indebida eventos confirmados que afecten datos del Cliente."),
        ("Subencargados", "Imponer obligaciones equivalentes e informar una lista vigente y cambios relevantes."),
        ("Cierre", "Apoyar retorno, eliminación o anonimización según el acuerdo y obligaciones legales."),
    ], [3000, 6360])
    add_heading(doc, "7. Obligaciones del Cliente")
    add_matrix(doc, ("Obligación", "Aplicación"), [
        ("Base jurídica", "Contar con autorización u otra base válida e informar a los titulares."),
        ("Instrucciones", "Impartir instrucciones lícitas, claras y documentadas."),
        ("Accesos", "Administrar usuarios, roles y bajas oportunamente."),
        ("Minimización", "No cargar datos excesivos, ilícitos o ajenos al objeto."),
        ("Derechos", "Atender de fondo solicitudes y decisiones frente a titulares y autoridades."),
    ], [3000, 6360])
    add_heading(doc, "8. Subencargados y flujos internacionales")
    add_matrix(doc, ("Proveedor", "Servicio", "Región / país", "Estado"), [
        ("Render", "Alojamiento y operación", "________________", "Por validar"),
        ("Cloudflare R2", "Archivos", "________________", "Por validar"),
        ("Sentry", "Observabilidad", "________________", "Por validar"),
        ("Proveedor SMTP", "Correo", "________________", "Pendiente"),
    ], [1900, 2500, 2300, 2660])
    add_heading(doc, "9. Incidentes")
    add_para(doc, "La notificación incluirá, cuando esté disponible, naturaleza, datos y titulares potencialmente afectados, medidas adoptadas, riesgos y contacto. El Cliente conserva las decisiones de comunicación a titulares y autoridades, con cooperación de Roustix.")
    add_heading(doc, "10. Retorno, eliminación y auditoría")
    add_key_values(doc, [
        ("Plazo de exportación", "________ días desde la terminación"),
        ("Eliminación de copias activas", "________ días después de la entrega / vencimiento"),
        ("Rotación residual de respaldos", "________ días"),
        ("Evidencia de cierre", "Acta y constancia de exportación, conservación o eliminación"),
        ("Auditoría", "Evidencia razonable coordinada, protegiendo secretos y datos de terceros"),
    ])
    add_heading(doc, "11. Firmas")
    add_signature_table(doc, "Responsable · Cliente", "Encargados · Roustix")
    return save_document(doc, "04-ACUERDO-TRANSMISION-DATOS-ROUSTIX.docx", "Acuerdo de Transmisión de Datos Roustix", "Anexo de encargo de datos", legal=True)


def build_final_uat_report() -> Path:
    doc = start_document(
        "Informe Final y Aceptación UAT",
        "Resultados, hallazgos y decisión del piloto",
        "PIL-CIE-001",
        "Plantilla de cierre",
        "Consolidar evidencia objetiva de lo probado y registrar qué acepta la empresa, qué queda pendiente y cuál es la recomendación de salida.",
    )
    add_heading(doc, "1. Identificación")
    add_key_values(doc, [
        ("Empresa / tenant", "____________________________________________________________"),
        ("Periodo del piloto", "Desde ____ / ____ / ______ hasta ____ / ____ / ______"),
        ("Plan y módulos", "____________________________________________________________"),
        ("Usuarios participantes", "____________________________________________________________"),
        ("Responsable del Cliente", "____________________________________________________________"),
        ("Responsable Roustix", "____________________________________________________________"),
    ])
    add_heading(doc, "2. Resumen ejecutivo")
    add_key_values(doc, [
        ("Objetivo evaluado", "____________________________________________________________"),
        ("Resultado general", "☐ Aprobado  ☐ Aprobado con pendientes  ☐ No aprobado"),
        ("Recomendación", "☐ Continuar a contrato  ☐ Extender piloto  ☐ Cerrar sin continuidad"),
        ("Observación principal", "____________________________________________________________"),
    ])
    add_heading(doc, "3. Matriz de aceptación")
    rows = [
        ("Acceso, recuperación y roles", "", "", ""),
        ("Empresa, sede, áreas y usuarios", "", "", ""),
        ("Activos, ficha y hoja de vida", "", "", ""),
        ("Incidencia → OT → jornadas → cierre", "", "", ""),
        ("Repuestos y costos", "", "", ""),
        ("PDF, Excel y documentos", "", "", ""),
        ("Almacenamiento y límites", "", "", ""),
        ("Auditoría, soporte y aislamiento", "", "", ""),
    ]
    add_matrix(doc, ("Escenario", "Resultado", "Evidencia", "Observación"), rows, [3300, 1500, 2200, 2360])
    add_heading(doc, "4. Hallazgos y acciones")
    add_matrix(doc, ("ID", "Severidad", "Hallazgo", "Responsable / fecha", "Estado"), [("", "", "", "", "") for _ in range(6)], [700, 1300, 3500, 2400, 1460])
    add_heading(doc, "5. Métricas y uso")
    add_key_values(doc, [
        ("Usuarios activos / concurrentes", "____________________________________________________________"),
        ("Activos e incidencias registrados", "____________________________________________________________"),
        ("OT y jornadas completadas", "____________________________________________________________"),
        ("Archivos / almacenamiento", "____________________________________________________________"),
        ("Solicitudes de soporte", "____________________________________________________________"),
        ("Incidentes críticos o altos", "____________________________________________________________"),
    ])
    add_heading(doc, "6. Aceptación")
    add_para(doc, "La aceptación confirma el resultado de las pruebas descritas; no renuncia a derechos ni convierte hallazgos pendientes en funciones aceptadas si no están expresamente identificados.")
    add_signature_table(doc, "Empresa piloto", "Roustix")
    return save_document(doc, "05-INFORME-FINAL-ACEPTACION-UAT-ROUSTIX.docx", "Informe Final y Aceptación UAT Roustix", "Resultados y aceptación del piloto")


def build_closure_act() -> Path:
    doc = start_document(
        "Acta de Cierre del Piloto",
        "Decisión final, obligaciones pendientes y transición",
        "PIL-CIE-002",
        "Plantilla de cierre",
        "Formalizar la terminación del piloto y evitar que queden accesos, cobros, datos o compromisos sin una decisión escrita.",
    )
    add_heading(doc, "1. Identificación")
    add_party_block(doc, include_two_operators=False)
    add_key_values(doc, [
        ("Acta de vinculación relacionada", "____________________________________________________________"),
        ("Inicio y terminación", "Desde ____ / ____ / ______ hasta ____ / ____ / ______"),
        ("Informe UAT", "Código / versión: ___________________________________________"),
    ])
    add_heading(doc, "2. Decisión de cierre")
    add_key_values(doc, [
        ("Resultado", "☐ Finalizado satisfactoriamente  ☐ Finalizado con pendientes  ☐ Terminado anticipadamente"),
        ("Continuidad", "☐ Contrato SaaS  ☐ Extensión documentada  ☐ Sin continuidad"),
        ("Fecha efectiva", "____ / ____ / ______  Hora: ______  Zona: America/Bogota"),
        ("Motivo / resumen", "____________________________________________________________"),
    ])
    add_heading(doc, "3. Obligaciones de salida")
    add_matrix(doc, ("Actividad", "Responsable", "Fecha límite", "Evidencia / estado"), [
        ("Exportación acordada", "", "", ""),
        ("Revocación de usuarios y sesiones", "", "", ""),
        ("Retiro de integraciones o credenciales", "", "", ""),
        ("Pago o saldo pendiente", "", "", ""),
        ("Conservación / eliminación", "", "", ""),
        ("Entrega de informe final", "", "", ""),
    ], [3500, 1900, 1800, 2160])
    add_heading(doc, "4. Pendientes posteriores")
    add_matrix(doc, ("ID", "Compromiso", "Responsable", "Fecha", "Criterio de cierre"), [("", "", "", "", "") for _ in range(5)], [700, 3100, 1900, 1500, 2160])
    add_heading(doc, "5. Declaración")
    add_para(doc, "Las partes dejan constancia del cierre en la fecha indicada. La confidencialidad, propiedad intelectual, tratamiento de datos, pagos y demás obligaciones que por su naturaleza sobrevivan continuarán según los documentos aplicables.")
    add_signature_table(doc)
    return save_document(doc, "06-ACTA-CIERRE-PILOTO-ROUSTIX.docx", "Acta de Cierre del Piloto Roustix", "Cierre y transición del piloto")


def build_data_exit_certificate() -> Path:
    doc = start_document(
        "Constancia de Exportación, Conservación o Eliminación",
        "Trazabilidad de los datos al terminar el piloto o servicio",
        "PIL-CIE-003",
        "Plantilla de cierre",
        "Registrar exactamente qué ocurrió con los datos y archivos del Cliente, sin declarar eliminaciones que no hayan sido verificadas.",
    )
    add_heading(doc, "1. Identificación")
    add_key_values(doc, [
        ("Empresa / tenant", "____________________________________________________________"),
        ("Solicitud / acta relacionada", "____________________________________________________________"),
        ("Responsable autorizado", "____________________________________________________________"),
        ("Fecha de corte", "____ / ____ / ______  Hora: ______  Zona: America/Bogota"),
    ])
    add_heading(doc, "2. Acción certificada")
    add_key_values(doc, [
        ("Tipo", "☐ Exportación  ☐ Conservación temporal  ☐ Eliminación  ☐ Combinación"),
        ("Alcance", "☐ BD  ☐ Archivos  ☐ Auditoría  ☐ Facturación  ☐ Soporte  ☐ Otros"),
        ("Periodo incluido", "____________________________________________________________"),
        ("Exclusiones", "____________________________________________________________"),
    ])
    add_heading(doc, "3. Inventario de exportación")
    add_matrix(doc, ("Archivo / conjunto", "Formato", "Registros / tamaño", "Hash / referencia", "Entrega"), [("", "", "", "", "") for _ in range(6)], [2700, 1100, 1900, 2100, 1560])
    add_heading(doc, "4. Conservación residual")
    add_matrix(doc, ("Categoría", "Motivo", "Ubicación / protección", "Fecha prevista de eliminación"), [
        ("Copias activas", "", "", ""),
        ("Respaldos", "Rotación y recuperación", "Bucket de recuperación restringido", ""),
        ("Auditoría / seguridad", "", "", ""),
        ("Legal / contable", "", "", ""),
    ], [2200, 2400, 3000, 1760])
    add_heading(doc, "5. Eliminación verificada")
    add_matrix(doc, ("Componente", "Método", "Fecha", "Resultado / evidencia"), [
        ("Base de datos activa", "", "", ""),
        ("Objetos / archivos activos", "", "", ""),
        ("Credenciales e integraciones", "", "", ""),
        ("Respaldos vencidos", "Rotación programada", "", ""),
    ], [2600, 2400, 1600, 2760])
    add_callout(doc, "Regla de evidencia", "Marque únicamente las acciones realmente ejecutadas. Si existen respaldos residuales o conservación legal, no declare eliminación total; identifique plazo, control de acceso y fecha prevista.", fill="FFF8E8", accent=GOLD)
    add_heading(doc, "6. Aceptación de entrega")
    add_signature_table(doc, "Cliente que recibe / valida", "Roustix que ejecuta / certifica")
    return save_document(doc, "07-CONSTANCIA-DATOS-SALIDA-ROUSTIX.docx", "Constancia de Datos de Salida Roustix", "Exportación, conservación o eliminación")


def build_saas_contract() -> Path:
    doc = start_document(
        "Contrato de Servicio SaaS",
        "Continuidad comercial posterior al piloto",
        "SAA-LEG-001",
        "Borrador sujeto a revisión jurídica",
        "Formalizar la continuidad del Cliente después del piloto. Solo debe firmarse tras completar datos, condiciones económicas, anexos y revisión jurídica, contable y tributaria.",
    )
    add_callout(doc, "No usar sin revisión", "Este modelo no fija todavía topes de responsabilidad, jurisdicción, impuestos ni SLA definitivo. Esos puntos requieren decisión de los socios y asesoría profesional antes de ofrecer el servicio.", fill="FCEDED", accent=RED)
    add_heading(doc, "1. Partes")
    add_party_block(doc)
    sections = [
        ("2. Objeto", "La parte prestadora habilita acceso empresarial no exclusivo a Roustix, bajo modalidad SaaS, para los módulos, usuarios, sedes, almacenamiento y período definidos en la orden o anexo comercial."),
        ("3. Documentos del contrato", "Integran el contrato la orden comercial, anexos de módulos y precios, política y acuerdo de datos, guía operativa, SLA cuando exista y modificaciones firmadas."),
        ("4. Vigencia y renovación", "La vigencia inicial, fecha de inicio, renovación, preaviso y permanencia se completan en la orden comercial. La continuidad no se presume cuando exista una condición pendiente."),
        ("5. Precio, impuestos y pago", "El Cliente pagará los valores expresamente pactados. Facturación, impuestos, vencimiento, mora, ajustes y medios de pago deben revisarse conforme a la situación legal y tributaria de la parte prestadora."),
        ("6. Acceso y uso", "El Cliente administrará usuarios y roles, protegerá credenciales, usará la plataforma de forma lícita y no intentará eludir controles, acceder a otros tenants, introducir código malicioso o explotar la propiedad intelectual."),
        ("7. Soporte y niveles de servicio", "Canales, horario, severidades, mantenimiento, objetivos de respuesta y disponibilidad se definen en la orden y, si aplica, en un SLA. La guía operativa regula el procedimiento, no compensaciones."),
        ("8. Seguridad y continuidad", "Roustix mantendrá controles razonables de autenticación, permisos, aislamiento, cifrado en tránsito, auditoría, almacenamiento persistente, monitoreo, respaldo y recuperación. El Cliente aplicará controles en sus dispositivos y personal."),
        ("9. Datos personales", "Las partes cumplirán los instrumentos de privacidad y transmisión. El Cliente responde por la licitud y minimización de los datos que incorpora; Roustix seguirá instrucciones documentadas como encargado cuando corresponda."),
        ("10. Confidencialidad", "Cada parte protegerá información no pública con medidas razonables, limitará su uso al contrato y notificará incidentes relevantes. La obligación sobrevivirá mientras la información conserve carácter confidencial."),
        ("11. Propiedad intelectual", "Roustix y sus componentes permanecen en cabeza de sus titulares. El Cliente conserva sus datos y contenidos. Ninguna disposición transfiere marcas, código o documentación salvo licencia expresa."),
        ("12. Subcontratación", "La parte prestadora podrá usar proveedores necesarios de infraestructura, almacenamiento, correo y observabilidad, manteniendo obligaciones contractuales y de datos aplicables."),
        ("13. Suspensión", "Podrá suspenderse acceso para contener riesgos, cumplir autoridad, atender uso ilícito, incumplimiento material o mora pactada, procurando proporcionalidad, aviso y oportunidad de subsanar cuando sea viable."),
        ("14. Terminación y salida", "Al terminar se revocarán accesos y se aplicarán exportación, conservación y eliminación conforme a la orden, el acuerdo de datos y la constancia de salida. Las obligaciones sobrevivientes continuarán."),
        ("15. Garantías y responsabilidad", "Las garantías, exclusiones permitidas, topes, indemnidades y seguros se completarán tras análisis jurídico y comercial. Nada excluirá responsabilidades que legalmente no puedan limitarse."),
        ("16. Fuerza mayor", "Ninguna parte responderá por incumplimientos causados exclusivamente por eventos irresistibles e imprevisibles en los términos legales, siempre que mitigue y comunique oportunamente."),
        ("17. Notificaciones", "Las direcciones contractuales y canales autorizados constarán en la carátula. Cambios deberán notificarse por un medio verificable."),
        ("18. Ley y controversias", "Ley aplicable, jurisdicción, negociación directa, conciliación u otro mecanismo se definirán con asesoría jurídica antes de firmar."),
        ("19. Integridad y cambios", "El contrato y anexos contienen el acuerdo aplicable. Las modificaciones requieren forma verificable y aceptación de representantes autorizados."),
    ]
    for heading, text in sections:
        add_heading(doc, heading)
        add_para(doc, text)
    add_heading(doc, "20. Orden comercial")
    add_key_values(doc, [
        ("Plan", "☐ Start  ☐ Business  ☐ Enterprise  ☐ Personalizado"),
        ("Precio / impuestos", "$ __________________ COP  ·  Impuestos: __________________"),
        ("Usuarios / sedes / módulos", "____________________________________________________________"),
        ("Almacenamiento / add-ons", "____________________________________________________________"),
        ("Inicio / vigencia / renovación", "____________________________________________________________"),
        ("Facturación / vencimiento", "____________________________________________________________"),
        ("Soporte / SLA", "____________________________________________________________"),
        ("Preaviso de terminación", "________ días"),
        ("Exportación / conservación", "____________________________________________________________"),
    ])
    add_heading(doc, "21. Aprobaciones pendientes")
    add_matrix(doc, ("Tema", "Decisión", "Revisión"), [
        ("Identidad jurídica de la parte prestadora", "", "Jurídica"),
        ("Facturación e impuestos", "", "Contable / tributaria"),
        ("Responsabilidad y seguros", "", "Jurídica / comercial"),
        ("SLA y compensaciones", "", "Operación / jurídica"),
        ("Ley y controversias", "", "Jurídica"),
        ("Conservación y eliminación", "", "Datos / jurídica"),
    ], [3600, 3600, 2160])
    add_heading(doc, "22. Firmas")
    add_signature_table(doc, "Cliente", "Prestador de Roustix")
    return save_document(doc, "08-CONTRATO-SAAS-ROUSTIX.docx", "Contrato de Servicio SaaS Roustix", "Continuidad posterior al piloto", legal=True)


def build_service_and_plan_act() -> Path:
    doc = start_document(
        "Acta de Servicio y Selección de Plan",
        "Activación comercial, alcance contratado y responsables",
        "SAA-ACT-001",
        "Plantilla para completar y firmar",
        "Dejar constancia del servicio Roustix que se activa para una empresa, el plan seleccionado, sus límites, complementos, valor, vigencia y responsables.",
        version="1.1",
    )
    add_callout(
        doc,
        "Documento complementario",
        "Esta acta concreta la orden de servicio. Debe leerse junto con el contrato SaaS o los términos del piloto, la Guía Operativa y los documentos de privacidad y transmisión aplicables.",
        fill="FFF8E8",
        accent=GOLD,
    )

    add_heading(doc, "1. Identificación del servicio")
    add_key_values(doc, [
        ("Empresa / razón social", "____________________________________________________________"),
        ("NIT / identificación", "____________________________________________________________"),
        ("Representante autorizado", "____________________________________________________________"),
        ("Contacto administrativo", "____________________________________________________________"),
        ("Contacto de facturación", "____________________________________________________________"),
        ("Tenant / slug", "____________________________________________________________"),
        ("Sector / vertical", "☐ Manufactura  ☐ Educación  ☐ Salud  ☐ Hotelería  ☐ Transporte  ☐ Otro"),
        ("Modalidad", "☐ Piloto controlado  ☐ Servicio SaaS posterior al piloto"),
    ])

    add_heading(doc, "2. Planes comerciales Roustix")
    add_matrix(doc, ("Plan", "Mensualidad", "Usuarios", "Sedes", "Módulos principales", "Almacenamiento", "Selección"), [
        ("Start", "$1.000.000 COP", "20", "1", "1 a elegir", "1 GB", "☐"),
        ("Business", "$1.500.000 COP", "50", "3", "Hasta 2", "5 GB", "☐"),
        ("Enterprise", "Desde $2.500.000 COP", "Personalizado", "Personalizadas", "Todos / según alcance", "20 GB ampliables", "☐"),
    ], [1200, 1800, 1050, 950, 1850, 1500, 1010], size=8.7)
    add_para(doc, "Enterprise se cotiza de acuerdo con el alcance definitivo. Los valores, impuestos, descuentos o condiciones especiales aplicables deben quedar diligenciados en la sección económica de esta acta.")

    add_heading(doc, "3. Plan y alcance seleccionado")
    add_key_values(doc, [
        ("Plan seleccionado", "☐ Start  ☐ Business  ☐ Enterprise  ☐ Personalizado"),
        ("Número de usuarios habilitados", "________________"),
        ("Número de sedes habilitadas", "________________"),
        ("Cuota base de almacenamiento", "________________ GB"),
        ("Módulo principal 1", "____________________________________________________________"),
        ("Módulo principal 2", "____________________________________________________________"),
        ("Otros módulos Enterprise", "____________________________________________________________"),
        ("Fecha prevista de activación", "____ / ____ / ______"),
    ])

    add_heading(doc, "4. Capacidades incluidas")
    add_matrix(doc, ("Capacidad", "Condición del servicio"), [
        ("Activos", "Ilimitados dentro del uso razonable, módulos habilitados y capacidad técnica acordada."),
        ("Seguridad", "HTTPS, autenticación, permisos por rol, aislamiento entre empresas y auditoría."),
        ("Archivos", "Almacenamiento persistente sujeto a la cuota efectiva del plan y complementos."),
        ("Respaldo", "Respaldos y recuperación conforme a la Guía Operativa y condiciones vigentes."),
        ("Actualizaciones", "Correcciones y mejoras generales de la plataforma durante la vigencia."),
        ("Soporte", "Email para Start; chat para Business; dedicado y SLA acordado para Enterprise."),
        ("Onboarding inicial", "Activación del tenant y configuración remota guiada del alcance contratado."),
        ("Capacitación inicial", "Sesión remota para administradores y responsables, según la propuesta aceptada."),
        ("Seguimiento de adopción", "Revisión inicial de uso y hallazgos dentro del periodo acordado."),
        ("Acceso multidispositivo", "Mediante navegador compatible y cuentas personales autorizadas."),
    ], [2600, 6760])

    add_heading(doc, "5. Complementos recurrentes")
    add_matrix(doc, ("Complemento", "SKU", "Valor mensual", "Cantidad", "Seleccionado"), [
        ("Almacenamiento +2 GB", "ADD-STG-2G", "$100.000 COP", "______", "☐"),
        ("Almacenamiento +5 GB", "ADD-STG-5G", "$220.000 COP", "______", "☐"),
        ("Almacenamiento +10 GB", "ADD-STG-10G", "$400.000 COP", "______", "☐"),
        ("Usuario adicional", "ADD-USR-COT", "$40.000 COP / usuario", "______", "☐"),
        ("Soporte prioritario", "ADD-SUP-PRI", "$350.000 COP", "______", "☐"),
    ], [2600, 1600, 2000, 1300, 1860], size=8.9)
    add_para(doc, "La activación de complementos sujetos a pago se realiza después de confirmar la condición comercial acordada. El retiro de capacidad no elimina automáticamente archivos existentes.")

    add_heading(doc, "6. Servicios profesionales opcionales")
    add_callout(
        doc,
        "Valores comerciales sugeridos",
        "Los siguientes valores son referencias para estructurar la cotización. No constituyen una tarifa definitiva hasta que el servicio, alcance, impuestos, desplazamientos y entregables queden aceptados por escrito.",
        fill="FFF8E8",
        accent=GOLD,
    )
    add_matrix(doc, ("Servicio", "SKU", "Unidad / alcance de referencia", "Valor sugerido", "Selección"), [
        ("Capacitación remota adicional", "ADD-TRN-REM", "Sesión de hasta 2 horas", "$250.000 COP", "☐"),
        ("Capacitación o implementación en sitio", "ADD-ONB-SITE", "Jornada de hasta 8 horas en Bogotá", "$900.000 COP + gastos", "☐"),
        ("Migración inicial de activos", "ADD-MIG-500", "Hasta 500 registros en plantilla validada", "$600.000 COP", "☐"),
        ("Depuración y organización de datos", "ADD-DATA-CLE", "Paquete inicial; alcance por diagnóstico", "Desde $600.000 COP", "☐"),
        ("Consultoría de configuración", "ADD-CNS-4H", "Bloque remoto de hasta 4 horas", "$500.000 COP", "☐"),
        ("Formato PDF o informe personalizado", "ADD-RPT-COT", "Cada formato aprobado", "Desde $450.000 COP", "☐"),
        ("Integración API / webhooks", "ADD-INT-COT", "Por integración y alcance técnico", "Desde $1.500.000 COP", "☐"),
        ("Desarrollo o personalización especial", "ADD-DEV-COT", "Bolsa o proyecto previamente estimado", "$180.000 COP / hora", "☐"),
    ], [2200, 1450, 2800, 1800, 1110], size=8.35)
    add_key_values(doc, [
        ("Servicio profesional seleccionado", "____________________________________________________________"),
        ("Alcance / entregable", "____________________________________________________________"),
        ("Fecha o plazo estimado", "____________________________________________________________"),
        ("Valor definitivo aprobado", "$ __________________________ COP"),
    ])

    add_heading(doc, "7. Condiciones económicas")
    add_key_values(doc, [
        ("Mensualidad base", "$ __________________________ COP"),
        ("Complementos mensuales", "$ __________________________ COP"),
        ("Servicios de pago único", "$ __________________________ COP"),
        ("Descuento / condición especial", "____________________________________________________________"),
        ("Subtotal", "$ __________________________ COP"),
        ("Impuestos aplicables", "____________________________________________________________"),
        ("Total periódico", "$ __________________________ COP"),
        ("Periodicidad", "☐ Mensual anticipada  ☐ Otra: ______________________________"),
        ("Medio de pago", "____________________________________________________________"),
        ("Fecha de vencimiento", "____________________________________________________________"),
    ])
    add_callout(doc, "Validación tributaria", "La parte prestadora debe validar facturación e impuestos aplicables según su situación jurídica y tributaria antes de emitir cobros o firmar la versión definitiva.", fill="FFF8E8", accent=GOLD)

    add_heading(doc, "8. Vigencia, renovación y terminación")
    add_key_values(doc, [
        ("Inicio del servicio", "____ / ____ / ______"),
        ("Vigencia inicial", "________ meses"),
        ("Renovación", "☐ Automática  ☐ Por acuerdo expreso  ☐ No aplica al piloto"),
        ("Preaviso de no renovación", "________ días calendario"),
        ("Plazo de exportación al terminar", "________ días"),
        ("Conservación residual de respaldos", "________ días / según política aprobada"),
    ])

    add_heading(doc, "9. Responsables y soporte")
    add_matrix(doc, ("Rol", "Nombre", "Correo / teléfono", "Responsabilidad"), [
        ("Administrador del Cliente", "", "", "Usuarios, permisos y configuración"),
        ("Responsable operativo", "", "", "Procesos, activos y mantenimiento"),
        ("Facturación del Cliente", "", "", "Pagos y documentos comerciales"),
        ("Onboarding Roustix", "", "", "Activación y acompañamiento"),
        ("Soporte Roustix", "", "", "Incidentes y escalamiento"),
    ], [1900, 1900, 2500, 3060])
    add_key_values(doc, [
        ("Canal temporal de soporte", "soporte.roustix@hotmail.com"),
        ("Horario acordado", "Lunes a viernes: ______ a ______ · America/Bogota"),
        ("Canal alterno", "____________________________________________________________"),
    ])

    add_heading(doc, "10. Gate de activación")
    add_matrix(doc, ("Verificación", "Estado", "Evidencia / observación"), [
        ("Empresa, plan y tenant correctamente identificados", "☐", ""),
        ("Usuarios, sedes y módulos dentro del alcance", "☐", ""),
        ("Administrador y responsables designados", "☐", ""),
        ("Términos o contrato aceptados", "☐", ""),
        ("Privacidad y transmisión entregadas / aceptadas", "☐", ""),
        ("Guía Operativa entregada", "☐", ""),
        ("Condición de pago validada", "☐", ""),
        ("Fecha y responsable de activación confirmados", "☐", ""),
    ], [5100, 900, 3360])

    add_heading(doc, "11. Documentos que integran el servicio")
    add_matrix(doc, ("Documento", "Código / versión", "Fecha / aceptación"), [
        ("Propuesta comercial", "", ""),
        ("Contrato SaaS o Términos del Piloto", "", ""),
        ("Política, aviso y autorización", "", ""),
        ("Acuerdo de transmisión de datos", "", ""),
        ("Guía Operativa", "", ""),
        ("Otros anexos", "", ""),
    ], [4200, 2400, 2760])

    add_heading(doc, "12. Aceptación y autorización de activación")
    add_para(doc, "Las partes confirman que el alcance, plan, capacidades, complementos, servicios profesionales, valores y responsables corresponden a lo acordado. La firma autoriza la activación del servicio en la fecha indicada, sin reemplazar las obligaciones de los demás documentos integrantes.")
    add_signature_table(doc, "Empresa Cliente", "Prestador / operadores de Roustix")
    return save_document(doc, "09-ACTA-SERVICIO-SELECCION-PLAN-ROUSTIX.docx", "Acta de Servicio y Selección de Plan Roustix", "Activación comercial y alcance del servicio")


def build_all() -> list[Path]:
    return [
        build_commercial_proposal(),
        build_pilot_terms(),
        build_privacy_package(),
        build_data_processing_agreement(),
        build_final_uat_report(),
        build_closure_act(),
        build_data_exit_certificate(),
        build_saas_contract(),
        build_service_and_plan_act(),
    ]


if __name__ == "__main__":
    for generated in build_all():
        print(generated)
