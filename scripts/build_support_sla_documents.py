"""Genera la Política de Soporte y el SLA corporativos de Roustix."""

from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from build_pilot_document_package import (
    BLACK,
    GOLD,
    NAVY,
    add_callout,
    add_heading,
    add_key_values,
    add_matrix,
    add_para,
    add_signature_table,
    set_run_font,
    start_document,
)


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = Path(
    os.environ.get(
        "ROUSTIX_SUPPORT_DOC_OUTPUT_DIR",
        ROOT / "docs" / "production-readiness" / "templates",
    )
)


def save_document(doc: Document, filename: str, title: str, subject: str) -> Path:
    props = doc.core_properties
    props.title = title
    props.subject = subject
    props.author = "Roustix"
    props.keywords = "Roustix, SaaS, soporte, SLA, niveles de servicio"
    props.comments = "Plantilla corporativa sujeta a aprobación jurídica, comercial y operativa."
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / filename
    doc.save(path)
    return path


def _numbering_instance(doc: Document, style_name: str) -> int:
    """Crea una instancia de numeración que reinicia en 1."""
    style = doc.styles[style_name]
    style_num_id = int(style._element.pPr.numPr.numId.val)
    numbering = doc.part.numbering_part.element
    source_num = numbering.xpath(f"./w:num[@w:numId='{style_num_id}']")[0]
    abstract_id = source_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    existing_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    new_id = max(existing_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return new_id


def add_list(doc: Document, items: list[str], *, numbered: bool = False) -> None:
    style = "List Number" if numbered else "List Bullet"
    num_id = _numbering_instance(doc, style) if numbered else None
    for item in items:
        paragraph = doc.add_paragraph(style=style)
        paragraph.paragraph_format.keep_together = True
        if num_id is not None:
            num_pr = paragraph._p.get_or_add_pPr().get_or_add_numPr()
            num_pr.get_or_add_ilvl().val = 0
            num_pr.get_or_add_numId().val = num_id
        run = paragraph.add_run(item)
        set_run_font(run, size=11, color=BLACK)


def build_support_policy() -> Path:
    doc = start_document(
        "Política de Soporte Roustix",
        "Canales, horarios, alcance y escalamiento del servicio de ayuda",
        "RTX-SUP-001",
        "Borrador para revisión y aprobación",
        "Definir cómo solicita y recibe ayuda un Cliente de Roustix, qué incluye el soporte y cómo se clasifican y escalan las solicitudes.",
        version="0.2.0",
    )
    add_callout(
        doc,
        "Alcance de esta política",
        "Esta política describe el servicio operativo de ayuda. Los objetivos contractuales de disponibilidad y tiempos de respuesta solo aplican cuando el Cliente tiene anexado y activado RTX-SLA-001.",
        fill="FFF8E8",
        accent=GOLD,
    )

    add_heading(doc, "1. Propósito")
    add_para(
        doc,
        "Definir canales, horarios, tipos de solicitud, inclusiones, exclusiones y escalamiento del soporte de Roustix para los planes comerciales vigentes.",
    )

    add_heading(doc, "2. Canales por plan")
    add_matrix(doc, ("Plan", "Canal principal", "Canal secundario", "Condición"), [
        ("Start", "Correo", "—", "Un hilo por incidente o consulta."),
        ("Business", "Chat, cuando esté habilitado", "Correo", "Misma cola de soporte."),
        ("Enterprise", "Canal dedicado, según Orden", "Correo / chat", "Puede incluir contacto nombrado."),
        ("ADD-SUP-PRI", "Priorización en cola", "Según plan base", "No crea por sí solo un SLA de disponibilidad."),
    ], [1500, 2300, 1800, 3760], size=8.9)
    doc.add_paragraph()
    add_key_values(doc, [
        ("Correo corporativo", "contacto@roustix.com"),
        ("Chat", "Dentro de la Plataforma para Business o superior, cuando esté habilitado"),
        ("Portal de soporte", "Funcionalidad futura; no disponible actualmente"),
    ])
    add_callout(
        doc,
        "Canales autorizados",
        "No se presta soporte por redes sociales ni por canales que no estén identificados en esta política, la propuesta, el contrato o la Orden de servicio.",
    )

    add_heading(doc, "3. Horarios")
    add_key_values(doc, [
        ("Días", "Lunes a viernes"),
        ("Horario", "8:00 a. m. a 5:00 p. m. · hora de Colombia (UTC−5)"),
        ("Excluidos", "Sábados, domingos y festivos colombianos"),
        ("Fuera de horario", "La solicitud se registra y se atiende el siguiente día hábil"),
        ("Condición especial", "Enterprise o un SLA anexado pueden pactar ventanas distintas en la Orden"),
    ])

    add_heading(doc, "4. Tipos de solicitudes")
    add_matrix(doc, ("Tipo", "Descripción", "Ejemplo"), [
        ("Incidente", "Error o degradación del servicio.", "No carga un módulo; error 500 recurrente."),
        ("Consulta", "Duda sobre el uso de la Plataforma.", "Cómo crear una orden de trabajo."),
        ("Configuración", "Ayuda para parametrizar lo incluido en el plan.", "Roles, sedes o catálogos."),
        ("Capacitación", "Orientación breve o sesión pactada.", "Recorrido de un módulo."),
        ("Mejora", "Sugerencia de producto.", "Nueva columna en un reporte."),
    ], [1800, 3560, 4000])
    add_para(doc, "Las mejoras no tienen plazo de implementación; se analizan y priorizan dentro del roadmap de producto.")

    add_heading(doc, "5. Lo que incluye")
    add_list(doc, [
        "Ayuda funcional sobre los módulos contratados.",
        "Diagnóstico y corrección de errores atribuibles a la Plataforma.",
        "Orientación básica de configuración dentro del alcance contratado.",
        "Información sobre el estado de incidentes conocidos.",
        "Guía para exportar datos o utilizar funciones documentadas.",
        "Onboarding inicial remoto conforme a la Orden de servicio.",
    ])

    add_heading(doc, "6. Lo que no incluye")
    add_para(doc, "Salvo cotización y aceptación expresa como servicio profesional:")
    add_list(doc, [
        "Desarrollo a medida o personalizaciones para un único Cliente.",
        "Cambios de producto solicitados por un solo tenant.",
        "Administración de infraestructura, red o dispositivos del Cliente.",
        "Soporte sobre equipos, VPN, firewalls o proveedor de Internet del Cliente.",
        "Recuperación de datos borrados por Usuarios fuera de la política de respaldos.",
        "Capacitación masiva o presencial no contratada.",
        "Integraciones API o webhooks no contratadas.",
        "Asesoría legal, contable o de cumplimiento del Cliente.",
    ])

    add_heading(doc, "7. Información mínima para abrir una solicitud")
    add_list(doc, [
        "Razón social y tenant.",
        "Usuario afectado.",
        "Descripción del problema y pasos para reproducirlo.",
        "Fecha y hora aproximada.",
        "Capturas o mensajes de error, sin contraseñas ni secretos.",
        "Impacto y confirmación de si bloquea la operación.",
    ], numbered=True)

    add_heading(doc, "8. Escalamiento interno")
    add_matrix(doc, ("Nivel", "Responsabilidad", "Criterio de escalamiento"), [
        ("Nivel 1 · Soporte funcional", "Clasificar, orientar y resolver consultas o configuraciones conocidas.", "No resuelto o sospecha de defecto."),
        ("Nivel 2 · Operación", "Diagnosticar infraestructura, logs, tenants e integraciones.", "Defecto de producto o cambio de código."),
        ("Desarrollo · Ingeniería", "Corregir defectos o evaluar cambios de producto.", "Según diagnóstico y priorización interna."),
    ], [2200, 4300, 2860], size=8.9)
    add_para(doc, "El Cliente no selecciona el nivel de atención; Roustix escala la solicitud según el diagnóstico.")

    add_heading(doc, "9. Severidad operativa")
    add_callout(
        doc,
        "Clasificación interna",
        "La severidad organiza la cola de atención. No constituye un compromiso de tiempo salvo que RTX-SLA-001 esté expresamente anexado y activado.",
        fill="FFF8E8",
        accent=GOLD,
    )
    add_matrix(doc, ("Severidad", "Ejemplo", "Tratamiento típico"), [
        ("P1 · Crítica", "Plataforma inaccesible para el tenant.", "Máxima prioridad en horario hábil."),
        ("P2 · Alta", "Función principal afectada sin alternativa.", "Prioridad alta."),
        ("P3 · Media", "Error parcial con alternativa disponible.", "Cola estándar."),
        ("P4 · Baja", "Consulta, mejora o defecto cosmético.", "Cola estándar o backlog."),
    ], [1800, 4100, 3460])

    add_heading(doc, "10. Relación con planes comerciales")
    add_matrix(doc, ("Plan", "Soporte incluido", "Documento aplicable"), [
        ("Start", "Correo", "Esta política."),
        ("Business", "Chat, cuando exista, y correo", "Esta política."),
        ("Enterprise", "Canal dedicado y SLA según contrato", "Esta política y RTX-SLA-001, si se anexa."),
    ], [1800, 3600, 3960])

    add_heading(doc, "11. Cambios")
    add_para(
        doc,
        "Roustix podrá actualizar esta política. Los cambios materiales se comunicarán a los clientes activos con preaviso razonable o en la renovación. La versión aplicable será la referenciada en el Contrato, la Orden o la publicada como vigente.",
    )

    add_heading(doc, "12. Control de cambios")
    add_matrix(doc, ("Versión", "Fecha", "Cambio"), [
        ("0.2.0", "2026-08-09", "Se establece contacto@roustix.com como canal corporativo y se retira el buzón temporal."),
        ("0.1.0", "2026-08-03", "Creación del borrador inicial."),
    ], [1500, 1800, 6060])

    return save_document(
        doc,
        "10-POLITICA-SOPORTE-ROUSTIX.docx",
        "Política de Soporte Roustix",
        "Canales, horarios, alcance y escalamiento del soporte",
    )


def build_sla() -> Path:
    doc = start_document(
        "Acuerdo de Nivel de Servicio (SLA)",
        "Objetivos propuestos de disponibilidad y respuesta",
        "RTX-SLA-001",
        "Borrador no vinculante · pendiente de confirmación operativa",
        "Definir objetivos de disponibilidad, prioridades, mantenimiento y exclusiones cuando el Contrato u Orden anexe expresamente este SLA.",
        version="0.1.0",
    )
    add_callout(
        doc,
        "Advertencia obligatoria",
        "Los porcentajes y tiempos son objetivos propuestos. Este documento no debe firmarse ni presentarse como compromiso hasta que Roustix confirme su capacidad operativa. Sin anexo expreso y activación escrita, no genera derechos ni créditos.",
        fill="FFF8E8",
        accent=GOLD,
    )

    add_heading(doc, "1. Objeto")
    add_para(
        doc,
        "Definir objetivos de disponibilidad, ventanas de soporte reforzado, prioridades de incidente, mantenimiento programado y exclusiones para el Cliente que tenga este SLA anexado.",
    )

    add_heading(doc, "2. Disponibilidad")
    add_heading(doc, "2.1 Objetivo mensual propuesto", level=2)
    add_matrix(doc, ("Nivel", "Objetivo mensual", "Uso sugerido"), [
        ("Estándar contractual", "99,5 %", "Primera etapa comercial o Business con SLA."),
        ("Reforzado", "99,9 %", "Solo tras validación operativa y cotización Enterprise."),
    ], [2600, 2000, 4760])
    add_callout(
        doc,
        "Fórmula de referencia",
        "Disponibilidad (%) = (minutos del mes − minutos de indisponibilidad atribuible) / minutos del mes × 100.",
    )
    add_list(doc, [
        "El mes corresponde al mes calendario en la zona horaria de Colombia.",
        "Indisponibilidad atribuible es la imposibilidad general de autenticarse o usar funciones críticas del tenant por causa imputable al Prestador, excluidas las causas de la sección 6.",
    ])

    add_heading(doc, "2.2 Créditos de servicio opcionales", level=2)
    add_callout(
        doc,
        "Decisión comercial pendiente",
        "Los créditos siguientes son una propuesta y no están vigentes hasta que se aprueben expresamente en la Orden o el Contrato.",
        fill="FFF8E8",
        accent=GOLD,
    )
    add_matrix(doc, ("Disponibilidad medida", "Crédito sugerido"), [
        ("Inferior al objetivo y desde 99,0 %", "5 % de la mensualidad del mes afectado"),
        ("Inferior a 99,0 % y desde 95,0 %", "10 % de la mensualidad del mes afectado"),
        ("Inferior a 95,0 %", "15 % de la mensualidad del mes afectado"),
    ], [4500, 4860])
    add_list(doc, [
        "El crédito sería el único remedio por incumplimiento de disponibilidad, salvo dolo.",
        "El tope mensual propuesto sería 15 % de la mensualidad del mes afectado.",
        "El Cliente debería solicitarlo dentro de los 30 días siguientes al mes afectado.",
        "No habría reembolso en efectivo; el crédito se aplicaría a facturas futuras.",
    ], numbered=True)

    add_heading(doc, "3. Horarios de soporte bajo SLA")
    add_key_values(doc, [
        ("Días", "Lunes a viernes"),
        ("Horario", "8:00 a. m. a 5:00 p. m. · hora de Colombia"),
        ("Canales", "Los de RTX-SUP-001 y canal dedicado si aplica"),
        ("P1 fuera de horario", "PENDIENTE: definir si existe guardia, mejor esfuerzo o no cobertura"),
    ])

    add_heading(doc, "4. Prioridades y tiempos objetivo")
    add_callout(
        doc,
        "Interpretación",
        "Los tiempos corresponden al primer contacto humano cualificado, no a la resolución, salvo estipulación expresa en la Orden.",
    )
    add_matrix(doc, ("Prioridad", "Ejemplo", "Respuesta objetivo"), [
        ("Crítica · P1", "Plataforma caída o tenant inaccesible.", "1 hora en el horario cubierto."),
        ("Alta · P2", "Función principal afectada.", "4 horas hábiles."),
        ("Media · P3", "Error parcial con alternativa disponible.", "8 horas hábiles."),
        ("Baja · P4", "Consulta o mejora.", "1 día hábil."),
    ], [1800, 4300, 3260])
    add_para(doc, "La resolución depende de la causa, la posibilidad de reproducir el incidente y su complejidad.")

    add_heading(doc, "5. Mantenimiento programado")
    add_list(doc, [
        "Roustix podrá realizar mantenimientos programados en ventanas de bajo uso.",
        "La ventana preferente propuesta es domingos de 00:00 a 06:00, hora de Colombia, u otra comunicada.",
        "Se procurará notificar con al menos 48 horas cuando sea razonablemente posible.",
        "El mantenimiento programado notificado no cuenta como indisponibilidad atribuible.",
        "Los mantenimientos de emergencia por seguridad pueden ejecutarse sin preaviso completo; se informarán tan pronto como sea viable.",
    ], numbered=True)

    add_heading(doc, "6. Exclusiones")
    add_para(doc, "No cuentan como indisponibilidad atribuible ni generan crédito:")
    add_list(doc, [
        "Fallas de Internet, DNS o equipos del Cliente.",
        "Navegadores no soportados o software del Cliente.",
        "Uso indebido, configuración errónea o acciones de Usuarios del Cliente.",
        "Incidentes de terceros fuera del control razonable de Roustix, siempre que se realicen acciones razonables de mitigación.",
        "Fuerza mayor.",
        "Suspensión legítima por mora, seguridad o uso ilícito.",
        "Características beta o preview expresamente identificadas.",
        "Periodos en los que el Cliente no coopere con el diagnóstico.",
    ])

    add_heading(doc, "7. Medición y reportes")
    add_matrix(doc, ("Elemento", "Definición / estado"), [
        ("Fuente de medición", "PENDIENTE: confirmar monitoreo propio, proveedor o combinación autorizada."),
        ("Reporte al Cliente", "Bajo demanda o mensual para Enterprise, según la Orden."),
        ("Discrepancias", "Se revisan con logs y evidencias disponibles de ambas partes."),
    ], [2700, 6660])

    add_heading(doc, "8. Relación con otros documentos")
    add_matrix(doc, ("Documento", "Relación"), [
        ("RTX-LEGAL-002", "Este SLA solo obliga cuando el Contrato u Orden lo anexa expresamente."),
        ("RTX-SUP-001", "Define canales y alcance; este SLA añade objetivos de tiempo y disponibilidad."),
        ("COM-01", "La referencia a SLA según contrato exige este documento u otro anexo pactado."),
    ], [2400, 6960])

    add_heading(doc, "9. Activación")
    add_para(doc, "El SLA queda activado únicamente cuando se cumplen simultáneamente estas condiciones:")
    add_list(doc, [
        "El Contrato o la Orden indican expresamente «SLA anexado: Sí».",
        "Ambas partes aceptan esta versión.",
        "Roustix confirma internamente que los objetivos son operativamente sostenibles.",
    ], numbered=True)
    add_key_values(doc, [
        ("SLA anexado", "☐ Sí  ☐ No"),
        ("Nivel seleccionado", "☐ Estándar 99,5 %  ☐ Reforzado 99,9 %  ☐ Otro: __________________"),
        ("Créditos aprobados", "☐ No aplican  ☐ Aplican según sección 2.2  ☐ Condición especial anexada"),
        ("P1 fuera de horario", "____________________________________________________________"),
        ("Fuente de medición", "____________________________________________________________"),
        ("Confirmación operativa Roustix", "Nombre / cargo / fecha: __________________________________________"),
    ])

    add_heading(doc, "10. Control de cambios")
    add_matrix(doc, ("Versión", "Fecha", "Cambio"), [
        ("0.1.0", "2026-08-03", "Borrador inicial con objetivos no vinculantes."),
    ], [1500, 1800, 6060])

    doc.add_page_break()
    add_heading(doc, "11. Aceptación del anexo")
    add_callout(
        doc,
        "Gate de firma",
        "No firmar mientras existan decisiones pendientes sin completar o la confirmación operativa de Roustix no esté diligenciada.",
        fill="FFF8E8",
        accent=GOLD,
    )
    add_signature_table(doc, "Empresa Cliente", "Prestador Roustix")

    return save_document(
        doc,
        "11-ACUERDO-NIVEL-SERVICIO-SLA-ROUSTIX.docx",
        "Acuerdo de Nivel de Servicio Roustix",
        "Objetivos propuestos de disponibilidad, respuesta y mantenimiento",
    )


def main() -> None:
    for path in (build_support_policy(), build_sla()):
        print(path)


if __name__ == "__main__":
    main()
