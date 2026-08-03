"""Genera la Guia Operativa para empresas piloto de Roustix."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from build_pilot_linkage_act import (
    BLACK,
    BLUE,
    CORPORATE_LOGO,
    GOLD,
    GRAY,
    GREEN,
    LIGHT_BLUE,
    NAVY,
    PALE_BLUE,
    RED,
    add_callout,
    add_cell_text,
    add_heading,
    add_word_field,
    configure_styles,
    enable_field_updates,
    set_cell_shading,
    set_run_font,
    set_table_geometry,
    style_table,
)


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "production-readiness" / "templates" / "GUIA-OPERATIVA-PILOTO-ROUSTIX.docx"


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("ROUSTIX  ·  GUÍA OPERATIVA DEL PILOTO")
    set_run_font(r, size=8.5, color=GRAY, bold=True)

    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    left, right = table.rows[0].cells
    add_cell_text(left, "Guía Operativa del Piloto · Versión 1.0", color=GRAY, size=8.5)
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Página ")
    set_run_font(r, size=8.5, color=GRAY)
    add_word_field(p, "PAGE")
    r = p.add_run(" de ")
    set_run_font(r, size=8.5, color=GRAY)
    add_word_field(p, "NUMPAGES")
    set_table_geometry(table, [6500, 2860], indent=0)


def add_title_block(doc: Document) -> None:
    if not CORPORATE_LOGO.is_file():
        raise FileNotFoundError(f"No se encontró el logotipo corporativo: {CORPORATE_LOGO}")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(5)
    shape = p.add_run().add_picture(str(CORPORATE_LOGO), width=Inches(2.6))
    shape._inline.docPr.set("title", "Roustix")
    shape._inline.docPr.set("descr", "Logotipo corporativo de Roustix")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Guía Operativa del Piloto")
    set_run_font(r, size=26, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("Acceso, exportación, respaldo, soporte y reporte de incidentes")
    set_run_font(r, size=12.5, color=GRAY)

    table = doc.add_table(rows=3, cols=4)
    rows = [
        ("Código", "PIL-GUI-001", "Versión", "1.0"),
        ("Empresa", "____________________________", "Tenant", "____________________________"),
        ("Entregada el", "____ / ____ / ______", "Responsable", "____________________________"),
    ]
    for row, values in zip(table.rows, rows):
        for index, value in enumerate(values):
            add_cell_text(row.cells[index], value, bold=index % 2 == 0, color=NAVY if index % 2 == 0 else BLACK)
            if index % 2 == 0:
                set_cell_shading(row.cells[index], LIGHT_BLUE)
    set_table_geometry(table, [1300, 3000, 1300, 3760])
    style_table(table, header=False, label_columns={0, 2})

    doc.add_paragraph()
    add_callout(
        doc,
        "Alcance de esta guía",
        "Este documento explica cómo operar durante el piloto y cómo pedir ayuda. No reemplaza los términos del servicio, la política de tratamiento de datos, el aviso de privacidad ni el acuerdo de transmisión que resulten aplicables.",
        fill="FFF8E8",
        accent=GOLD,
    )


def add_procedure_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    for cell, text in zip(table.rows[0].cells, ("Paso", "Acción", "Resultado esperado")):
        add_cell_text(cell, text, bold=True, color="FFFFFF")
    for step, action, result in rows:
        cells = table.add_row().cells
        add_cell_text(cells[0], step, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_cell_text(cells[1], action)
        add_cell_text(cells[2], result)
    set_table_geometry(table, [720, 4740, 3900])
    style_table(table, header=True)


def add_roles_and_access(doc: Document) -> None:
    add_heading(doc, "1. Acceso, roles y confidencialidad")
    p = doc.add_paragraph(
        "Cada cuenta es personal. La empresa piloto administra quién accede, qué rol recibe y cuándo debe desactivarse. Roustix aplica los permisos configurados para el tenant, pero la empresa conserva la responsabilidad de asignarlos y revisarlos."
    )
    p.paragraph_format.keep_together = True

    table = doc.add_table(rows=1, cols=3)
    for cell, text in zip(table.rows[0].cells, ("Responsable", "Responsabilidades mínimas", "Evidencia")):
        add_cell_text(cell, text, bold=True, color="FFFFFF")
    data = [
        ("Administrador de empresa", "Crear, modificar y desactivar usuarios; asignar el menor privilegio necesario; revisar accesos.", "Listado de usuarios y auditoría"),
        ("Usuario", "Proteger su contraseña; no compartir la sesión; cerrar sesión en equipos compartidos; reportar actividad extraña.", "Aceptación de esta guía"),
        ("Supervisor", "Validar cierres, cambios sensibles y registros ejecutados por su equipo.", "Historial y auditoría"),
        ("Roustix", "Mantener controles de acceso, aislamiento entre empresas, trazabilidad y soporte de plataforma.", "Logs y panel operativo"),
    ]
    for values in data:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            add_cell_text(cell, value)
    set_table_geometry(table, [2100, 4740, 2520])
    style_table(table, header=True)

    add_callout(
        doc,
        "Reglas obligatorias",
        "No enviar contraseñas, códigos de verificación, cookies, API keys ni secretos por correo o capturas. Ante pérdida del dispositivo, sospecha de acceso o retiro de una persona, informe al administrador y solicite revocación inmediata de la cuenta o sesiones.",
        fill="FCEDED",
        accent=RED,
    )


def add_exports(doc: Document) -> None:
    add_heading(doc, "2. Exportación y conservación de información")
    p = doc.add_paragraph(
        "Las exportaciones disponibles dependen del módulo y del plan habilitado. Los botones PDF o Excel de cada vista exportan la información visible según sus filtros y permisos. Para una salida integral o un volumen que no pueda obtenerse desde la interfaz, el administrador debe solicitar asistencia a soporte."
    )
    p.paragraph_format.keep_together = True

    add_procedure_table(doc, [
        ("1", "Inicie sesión con un rol autorizado y abra la lista o registro requerido.", "Solo se muestran datos de la empresa y del alcance permitido."),
        ("2", "Aplique período, estado, activo, sede u otros filtros disponibles.", "La exportación respeta el conjunto filtrado cuando la función lo admite."),
        ("3", "Seleccione PDF o Excel y guarde el archivo en una ubicación controlada.", "Se obtiene una copia para consulta o entrega interna."),
        ("4", "Revise que el archivo incluya el período, empresa y registros esperados.", "La empresa confirma integridad antes de usar o compartir la copia."),
        ("5", "Si falta información o necesita una exportación integral, contacte soporte indicando tenant y alcance; no envíe datos sensibles innecesarios.", "La solicitud queda registrada y se define el método seguro de entrega."),
    ])

    add_callout(
        doc,
        "Custodia de las copias",
        "Una exportación deja de estar protegida por los controles de acceso de Roustix. La empresa debe almacenarla, compartirla y eliminarla conforme a sus propias políticas y a las reglas de tratamiento de datos aplicables.",
        fill=PALE_BLUE,
        accent=BLUE,
    )


def add_backup(doc: Document) -> None:
    add_heading(doc, "3. Respaldo, recuperación y continuidad")
    p = doc.add_paragraph(
        "Durante el piloto, Roustix ejecuta un respaldo diario de la base de datos y del inventario de archivos. La base y los archivos forman una única unidad lógica de recuperación. Los respaldos se validan mediante restauración técnica y simulacros documentados."
    )
    p.paragraph_format.keep_together = True

    table = doc.add_table(rows=1, cols=3)
    for cell, text in zip(table.rows[0].cells, ("Compromiso operativo del piloto", "Objetivo", "Aclaración")):
        add_cell_text(cell, text, bold=True, color="FFFFFF")
    rows = [
        ("Frecuencia de respaldo", "Diaria · 03:00 UTC", "La ventana puede modificarse por mantenimiento y queda monitoreada."),
        ("RPO objetivo", "Hasta 24 horas", "Máximo de información que podría requerir reconstrucción ante una pérdida total."),
        ("RTO de prueba", "Hasta 2 horas", "Tiempo objetivo para dejar un entorno recuperado utilizable durante un simulacro."),
        ("RTO de producción", "Hasta 4 horas", "Objetivo técnico de recuperación; no constituye todavía un SLA contractual."),
    ]
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            add_cell_text(cell, value)
    set_table_geometry(table, [2800, 1800, 4760])
    style_table(table, header=True)

    add_procedure_table(doc, [
        ("1", "El administrador reporta el evento y especifica qué información parece afectada.", "Se abre un caso con fecha, tenant y alcance inicial."),
        ("2", "Roustix confirma si corresponde restaurar, corregir datos o recuperar un archivo.", "Se evita sobrescribir producción sin diagnóstico y aprobación."),
        ("3", "La restauración se prepara primero en un destino aislado y se validan datos, acceso y una muestra de archivos.", "La recuperación se comprueba antes de cualquier cambio productivo."),
        ("4", "El responsable autorizado aprueba el retorno o corte a producción cuando resulte necesario.", "La decisión y la evidencia quedan registradas."),
    ])


def add_support(doc: Document) -> None:
    add_heading(doc, "4. Soporte y escalamiento")
    table = doc.add_table(rows=5, cols=2)
    values = [
        ("Primer nivel", "Administrador o supervisor de la empresa piloto"),
        ("Canal temporal Roustix", "soporte.roustix@hotmail.com"),
        ("Horario acordado", "Lunes a viernes: ______ a ______ · Zona America/Bogota"),
        ("Responsable Roustix", "____________________________________________________________"),
        ("Canal alterno / teléfono", "____________________________________________________________"),
    ]
    for row, pair in zip(table.rows, values):
        add_cell_text(row.cells[0], pair[0], bold=True, color=NAVY)
        add_cell_text(row.cells[1], pair[1])
        set_cell_shading(row.cells[0], LIGHT_BLUE)
    set_table_geometry(table, [2700, 6660])
    style_table(table, header=False, label_columns={0})

    table = doc.add_table(rows=1, cols=4)
    for cell, text in zip(table.rows[0].cells, ("Nivel", "Ejemplo", "Acción inicial", "Escalamiento")):
        add_cell_text(cell, text, bold=True, color="FFFFFF")
    rows = [
        ("P1 · Crítico", "Aplicación inaccesible, pérdida o exposición confirmada de datos.", "Reportar inmediatamente y conservar evidencia.", "Responsable Roustix + operación técnica."),
        ("P2 · Alto", "Flujo esencial bloqueado sin alternativa razonable.", "Reportar el mismo día con impacto y usuarios afectados.", "Soporte funcional y técnico."),
        ("P3 · Medio", "Error parcial con alternativa disponible.", "Registrar pasos, pantalla y resultado esperado.", "Soporte funcional."),
        ("P4 · Consulta", "Duda, mejora o acompañamiento.", "Enviar descripción y prioridad deseada.", "Onboarding o producto."),
    ]
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            add_cell_text(cell, value, size=9.2)
    set_table_geometry(table, [1200, 3000, 3000, 2160])
    style_table(table, header=True)


def add_incident_report(doc: Document) -> None:
    add_heading(doc, "5. Cómo reportar un incidente")
    add_procedure_table(doc, [
        ("1", "Confirme si el problema afecta a una persona, un rol o a toda la empresa.", "Se define el alcance sin repetir acciones peligrosas."),
        ("2", "Registre fecha y hora local, módulo, URL, acción realizada y mensaje exacto.", "El caso puede correlacionarse con los logs."),
        ("3", "Adjunte una captura saneada. Oculte datos personales, contraseñas, códigos y secretos.", "La evidencia es útil y segura."),
        ("4", "Indique impacto operativo y si existe una alternativa temporal.", "Soporte asigna la severidad correcta."),
        ("5", "Envíe el reporte al administrador y, cuando corresponda, al canal Roustix.", "Se genera seguimiento y cierre documentado."),
    ])

    add_heading(doc, "Formato rápido de reporte", level=2)
    table = doc.add_table(rows=8, cols=2)
    fields = [
        ("Empresa / tenant", ""),
        ("Reportado por / contacto", ""),
        ("Fecha y hora local", ""),
        ("Módulo, página o URL", ""),
        ("Descripción y pasos", ""),
        ("Mensaje de error", ""),
        ("Impacto / usuarios afectados", ""),
        ("Evidencia saneada", "☐ Captura  ☐ Archivo  ☐ ID de solicitud  ☐ No aplica"),
    ]
    for row, pair in zip(table.rows, fields):
        add_cell_text(row.cells[0], pair[0], bold=True, color=NAVY)
        add_cell_text(row.cells[1], pair[1] or "____________________________________________________________")
        set_cell_shading(row.cells[0], LIGHT_BLUE)
    set_table_geometry(table, [2800, 6560])
    style_table(table, header=False, label_columns={0})


def add_storage_and_exit(doc: Document) -> None:
    add_heading(doc, "6. Almacenamiento y límites")
    table = doc.add_table(rows=1, cols=4)
    for cell, text in zip(table.rows[0].cells, ("Plan", "Cuota base", "Alerta", "Al llegar al límite")):
        add_cell_text(cell, text, bold=True, color="FFFFFF")
    rows = [
        ("Start", "1 GB", "Desde 80 %", "Se bloquean nuevas cargas al 100 %."),
        ("Business", "5 GB", "Desde 80 %", "Puede liberar espacio, ampliar o cambiar de plan."),
        ("Enterprise", "20 GB ampliables", "Desde 80 %", "La ampliación se acuerda comercialmente."),
    ]
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            add_cell_text(cell, value)
    set_table_geometry(table, [1700, 1900, 1700, 4060])
    style_table(table, header=True)

    add_callout(
        doc,
        "Archivos al 100 %",
        "Las descargas y eliminaciones permanecen disponibles para liberar espacio. El piloto dispone de un add-on de +2 GB sujeto a aprobación y condiciones comerciales vigentes. La ampliación no sustituye las prácticas de archivo y depuración de la empresa.",
        fill=PALE_BLUE,
        accent=BLUE,
    )

    add_heading(doc, "7. Finalización del piloto y continuidad")
    p = doc.add_paragraph(
        "Antes de cerrar el piloto, la empresa y Roustix deben acordar por escrito la fecha de terminación, las exportaciones necesarias, el plazo de acceso posterior, la conservación o eliminación de archivos y cualquier migración. Los plazos definitivos se rigen por el acta, los términos y los documentos de tratamiento aprobados; esta guía no los sustituye."
    )
    p.paragraph_format.keep_together = True

    table = doc.add_table(rows=1, cols=3)
    for cell, text in zip(table.rows[0].cells, ("Actividad de cierre", "Responsable", "Evidencia")):
        add_cell_text(cell, text, bold=True, color="FFFFFF")
    rows = [
        ("Confirmar datos y archivos que deben exportarse", "Empresa piloto", "Solicitud aprobada"),
        ("Entregar o habilitar las exportaciones acordadas", "Roustix", "Archivos / acta de entrega"),
        ("Desactivar usuarios y accesos que no continúan", "Ambas partes", "Auditoría de accesos"),
        ("Aplicar conservación o eliminación según documentos vigentes", "Roustix", "Registro de cierre"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            add_cell_text(cell, value)
    set_table_geometry(table, [4500, 2100, 2760])
    style_table(table, header=True)


def add_acknowledgement(doc: Document) -> None:
    add_heading(doc, "8. Constancia de entrega y entendimiento")
    p = doc.add_paragraph(
        "Con la firma de esta constancia, la empresa confirma que recibió la guía, conoce el canal de soporte y comunicará a sus usuarios las responsabilidades de acceso y confidencialidad. La aceptación de privacidad y tratamiento se documenta separadamente con los instrumentos aplicables."
    )
    p.paragraph_format.keep_together = True

    table = doc.add_table(rows=6, cols=2)
    rows = [
        ("Empresa piloto", "Roustix"),
        ("Nombre: ______________________________", "Nombre: ______________________________"),
        ("Cargo: _______________________________", "Rol: __________________________________"),
        ("Firma: _______________________________", "Firma: _________________________________"),
        ("Fecha: ____ / ____ / ______", "Fecha: ____ / ____ / ______"),
        ("Versión recibida: ____________________", "Medio de entrega: ______________________"),
    ]
    for row, values in zip(table.rows, rows):
        for index, value in enumerate(values):
            add_cell_text(row.cells[index], value, bold=row._index == 0, color=NAVY if row._index == 0 else BLACK)
            if row._index == 0:
                set_cell_shading(row.cells[index], LIGHT_BLUE)
    set_table_geometry(table, [4680, 4680])
    style_table(table, header=False)


def build() -> Path:
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    enable_field_updates(doc)
    add_title_block(doc)
    add_roles_and_access(doc)
    add_exports(doc)
    add_backup(doc)
    add_support(doc)
    add_incident_report(doc)
    add_storage_and_exit(doc)
    add_acknowledgement(doc)

    props = doc.core_properties
    props.title = "Guía Operativa del Piloto Roustix"
    props.subject = "Acceso, exportación, respaldo, soporte y reporte de incidentes"
    props.author = "Roustix"
    props.keywords = "Roustix, piloto, acceso, exportación, respaldo, soporte, incidentes"
    props.comments = "Plantilla operativa para completar y entregar a cada empresa piloto."

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
